#!/usr/bin/env python3
"""
Synthèse Co-fabriquée par Conseil de LLMs
Orchestre Claude, Gemini et Codex pour produire une synthèse robuste.

Version 2.2 - Retry avec backoff (PRD-004)

Inspiré de Council (bacoco/Council-board-skill)
"""

import argparse
import asyncio
import json
import os
import subprocess
import sys
import time
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Any, Tuple, Callable, Awaitable, Protocol, runtime_checkable
from dataclasses import dataclass, field, asdict
from enum import Enum
import hashlib


# ══════════════════════════════════════════════════════════════════════════════
# PROTOCOLES POUR INJECTION DE DÉPENDANCES (PRD-015 Story 15.1)
# ══════════════════════════════════════════════════════════════════════════════

@runtime_checkable
class CLIExecutor(Protocol):
    """
    Protocole pour l'exécution de commandes CLI.

    PRD-015: Permet l'injection de dépendances pour les tests.
    """

    async def run(
        self,
        cmd: List[str],
        timeout: int = 300
    ) -> Tuple[str, str, int]:
        """
        Exécute une commande CLI.

        Args:
            cmd: Commande et arguments
            timeout: Timeout en secondes

        Returns:
            Tuple (stdout, stderr, returncode)
        """
        ...


class DefaultCLIExecutor:
    """
    Exécuteur CLI par défaut utilisant asyncio.create_subprocess_exec.

    PRD-015: Implémentation standard, peut être remplacée dans les tests.
    """

    async def run(
        self,
        cmd: List[str],
        timeout: int = 300
    ) -> Tuple[str, str, int]:
        """Exécute une commande via asyncio subprocess."""
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )

        try:
            stdout, stderr = await asyncio.wait_for(
                proc.communicate(),
                timeout=timeout
            )
            return (
                stdout.decode('utf-8', errors='replace'),
                stderr.decode('utf-8', errors='replace'),
                proc.returncode
            )
        except asyncio.TimeoutError:
            proc.kill()
            await proc.wait()
            raise

# Import du module de configuration (PRD-003)
try:
    from config import (
        ConfigLoader, ConfigError, DEFAULT_CONFIG,
        get_expert_roles, get_default_cadrage
    )
    CONFIG_MODULE_AVAILABLE = True
except ImportError:
    CONFIG_MODULE_AVAILABLE = False
    # Fallback si le module n'est pas disponible
    DEFAULT_CONFIG = {
        "max_rounds": 2,
        "timeout": 300,
        "enable_trail": True,
        "trail_dir": "synthese_trails",
        "models": ["claude", "gemini", "codex"],
        "convergence_threshold": 0.8,
        "fallback_to_available": True,
    }

# Import du module retry (PRD-004)
try:
    from retry import (
        retry, is_retryable, classify_error,
        RetryMetricsCollector, PrintRetryCallback,
        get_retry_metrics, reset_retry_metrics,
        DEFAULT_RETRY_CONFIG, RetryableError
    )
    RETRY_MODULE_AVAILABLE = True
except ImportError:
    RETRY_MODULE_AVAILABLE = False
    DEFAULT_RETRY_CONFIG = {
        "enabled": False,
        "max_attempts": 1,
    }

# Import du module pédagogique (PRD-005)
try:
    from pedagogy import (
        PedagogicPresenter, PedagogicStep,
        create_pedagogy, get_expert_explanation
    )
    PEDAGOGY_MODULE_AVAILABLE = True
except ImportError:
    PEDAGOGY_MODULE_AVAILABLE = False

# Import du module convergence (PRD-007)
try:
    from convergence import (
        calculate_convergence as calc_convergence_hybrid,
        interpret_convergence, get_keywords
    )
    CONVERGENCE_MODULE_AVAILABLE = True
except ImportError:
    CONVERGENCE_MODULE_AVAILABLE = False

# Import du module backends (PRD-008)
try:
    from backends import (
        BackendRegistry, LLMBackend, LLMResponse as BackendResponse,
        check_anthropic_availability, AnthropicBackend
    )
    BACKENDS_MODULE_AVAILABLE = True
except ImportError:
    BACKENDS_MODULE_AVAILABLE = False

# Import du module sanitize (PRD-008 Story 8.4)
try:
    from sanitize import (
        validate_and_sanitize, sanitize_prompt, sanitize_cadrage,
        validate_source_text, validate_cadrage as validate_cadrage_fn,
        ValidationError
    )
    SANITIZE_MODULE_AVAILABLE = True
except ImportError:
    SANITIZE_MODULE_AVAILABLE = False

# PRD-018: Import des extracteurs optionnels
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import requests
    from bs4 import BeautifulSoup
    URL_AVAILABLE = True
except ImportError:
    URL_AVAILABLE = False


# PRD-018: Dataclass Source pour multi-sources
@dataclass
class Source:
    """Représente une source de texte (PRD-018)."""
    name: str           # Nom de la source (fichier, URL)
    content: str        # Contenu textuel
    source_type: str    # "file", "pdf", "docx", "url", "text"
    metadata: Dict[str, Any] = field(default_factory=dict)


# PRD-019: Dataclass IndexedSource pour références
@dataclass
class IndexedSource:
    """
    Source avec index de paragraphes pour références [§N].

    PRD-019: Permet de créer des références vérifiables vers le texte source.
    """
    content: str
    paragraphs: List[str]
    index: Dict[int, Tuple[int, int]]  # para_num -> (start, end)
    source_name: str = "source"


# ══════════════════════════════════════════════════════════════════════════════
# INDEXATION ET RÉFÉRENCES (PRD-019)
# ══════════════════════════════════════════════════════════════════════════════

def index_source(text: str, source_name: str = "source") -> IndexedSource:
    """
    Indexe un texte source par paragraphes.

    PRD-019: Crée un mapping paragraphe -> position pour les références [§N].

    Args:
        text: Texte source à indexer
        source_name: Nom de la source (pour multi-sources)

    Returns:
        IndexedSource avec paragraphes numérotés
    """
    paragraphs = []
    index = {}
    current_pos = 0

    # Séparer par doubles sauts de ligne ou lignes vides
    raw_paras = re.split(r'\n\s*\n', text)

    for i, para in enumerate(raw_paras, 1):
        para = para.strip()
        if para:
            start = text.find(para, current_pos)
            end = start + len(para)
            paragraphs.append(para)
            index[i] = (start, end)
            current_pos = end

    return IndexedSource(
        content=text,
        paragraphs=paragraphs,
        index=index,
        source_name=source_name
    )


def format_indexed_for_prompt(indexed: IndexedSource) -> str:
    """
    Formate le texte source avec numéros de paragraphes.

    PRD-019: Prépare le texte pour le prompt de synthèse avec références.

    Args:
        indexed: Source indexée

    Returns:
        Texte formaté avec [§N] devant chaque paragraphe
    """
    lines = []
    for i, para in enumerate(indexed.paragraphs, 1):
        lines.append(f"[§{i}] {para}")
    return "\n\n".join(lines)


def format_multi_indexed_for_prompt(
    indexed_sources: List[Tuple[str, IndexedSource]]
) -> str:
    """
    Formate plusieurs sources indexées pour le prompt.

    PRD-019: Support multi-sources avec références [S:nom§N].

    Args:
        indexed_sources: Liste de (nom, IndexedSource)

    Returns:
        Texte formaté avec références multi-sources
    """
    parts = []
    for name, indexed in indexed_sources:
        header = f"═══ SOURCE: {name} ═══"
        formatted = []
        for i, para in enumerate(indexed.paragraphs, 1):
            formatted.append(f"[S:{name}§{i}] {para}")
        parts.append(f"{header}\n\n" + "\n\n".join(formatted))

    return "\n\n".join(parts)


def validate_references(
    synthesis: str,
    indexed: IndexedSource
) -> Tuple[bool, List[str], List[str]]:
    """
    Vérifie que les références existent dans la source.

    PRD-019: Détecte les références invalides ou inventées.

    Args:
        synthesis: Texte de synthèse avec références [§N]
        indexed: Source indexée

    Returns:
        Tuple (valid, errors, valid_refs)
    """
    errors = []
    valid_refs = []

    # Pattern pour références simples [§N] ou [§N.M]
    refs = re.findall(r'\[§(\d+(?:\.\d+)?)\]', synthesis)

    max_para = len(indexed.paragraphs)

    for ref in refs:
        para_num = int(ref.split('.')[0])
        if para_num < 1 or para_num > max_para:
            errors.append(f"[§{ref}] invalide (max: §{max_para})")
        else:
            valid_refs.append(f"§{ref}")

    return (len(errors) == 0, errors, valid_refs)


def validate_multi_references(
    synthesis: str,
    indexed_sources: Dict[str, IndexedSource]
) -> Tuple[bool, List[str], List[str]]:
    """
    Vérifie les références multi-sources [S:nom§N].

    PRD-019: Validation pour synthèses multi-sources.

    Args:
        synthesis: Texte de synthèse
        indexed_sources: Dict nom -> IndexedSource

    Returns:
        Tuple (valid, errors, valid_refs)
    """
    errors = []
    valid_refs = []

    # Pattern pour références multi-sources [S:nom§N]
    multi_refs = re.findall(r'\[S:([^\]§]+)§(\d+)\]', synthesis)

    for source_name, para_num_str in multi_refs:
        para_num = int(para_num_str)
        if source_name not in indexed_sources:
            errors.append(f"[S:{source_name}§{para_num}] source inconnue")
        else:
            max_para = len(indexed_sources[source_name].paragraphs)
            if para_num < 1 or para_num > max_para:
                errors.append(f"[S:{source_name}§{para_num}] invalide (max: §{max_para})")
            else:
                valid_refs.append(f"S:{source_name}§{para_num}")

    # Aussi valider les références simples [§N] si une seule source
    simple_refs = re.findall(r'\[§(\d+)\]', synthesis)
    if len(indexed_sources) == 1:
        indexed = list(indexed_sources.values())[0]
        max_para = len(indexed.paragraphs)
        for ref in simple_refs:
            para_num = int(ref)
            if para_num < 1 or para_num > max_para:
                errors.append(f"[§{ref}] invalide (max: §{max_para})")
            else:
                valid_refs.append(f"§{ref}")

    return (len(errors) == 0, errors, valid_refs)


def extract_reference_excerpts(
    valid_refs: List[str],
    indexed_sources: Dict[str, IndexedSource],
    max_excerpt_len: int = 50
) -> List[Dict[str, str]]:
    """
    Extrait les passages correspondant aux références.

    PRD-019: Génère la section Références avec extraits.

    Args:
        valid_refs: Liste des références valides
        indexed_sources: Dict nom -> IndexedSource
        max_excerpt_len: Longueur max des extraits

    Returns:
        Liste de {ref, excerpt}
    """
    excerpts = []

    for ref in valid_refs:
        if ref.startswith("S:"):
            # Multi-source: S:nom§N
            match = re.match(r'S:([^§]+)§(\d+)', ref)
            if match:
                source_name, para_num = match.groups()
                para_num = int(para_num)
                if source_name in indexed_sources:
                    indexed = indexed_sources[source_name]
                    if 0 < para_num <= len(indexed.paragraphs):
                        text = indexed.paragraphs[para_num - 1]
                        excerpt = text[:max_excerpt_len] + "..." if len(text) > max_excerpt_len else text
                        excerpts.append({"ref": f"[S:{source_name}§{para_num}]", "excerpt": excerpt})
        else:
            # Simple: §N
            para_num = int(ref.replace("§", ""))
            # Prendre la première source
            if indexed_sources:
                indexed = list(indexed_sources.values())[0]
                if 0 < para_num <= len(indexed.paragraphs):
                    text = indexed.paragraphs[para_num - 1]
                    excerpt = text[:max_excerpt_len] + "..." if len(text) > max_excerpt_len else text
                    excerpts.append({"ref": f"[§{para_num}]", "excerpt": excerpt})

    return excerpts


# ══════════════════════════════════════════════════════════════════════════════
# TEMPLATES ET CADRAGE (PRD-020)
# ══════════════════════════════════════════════════════════════════════════════

def resolve_cadrage(
    args: "argparse.Namespace",
    default_cadrage: Dict[str, str]
) -> Dict[str, str]:
    """
    Résout le cadrage selon la priorité : CLI > template > défaut.

    PRD-020: Support des templates de cadrage.

    Args:
        args: Arguments CLI parsés
        default_cadrage: Cadrage par défaut

    Returns:
        Cadrage résolu
    """
    # Si template spécifié, l'utiliser comme base
    template_name = getattr(args, 'template', None)
    if template_name:
        if template_name not in CADRAGE_TEMPLATES:
            raise ValueError(f"Template inconnu: {template_name}")
        cadrage = CADRAGE_TEMPLATES[template_name].copy()
        cadrage["_template"] = template_name  # Métadonnée
    else:
        cadrage = default_cadrage.copy()

    # Les options CLI explicites ont priorité sur le template
    if getattr(args, 'destinataire', None):
        cadrage["destinataire"] = args.destinataire
    if getattr(args, 'finalite', None):
        cadrage["finalite"] = args.finalite
    if getattr(args, 'longueur', None):
        cadrage["longueur"] = args.longueur
    if getattr(args, 'ton', None):
        cadrage["ton"] = args.ton
    if getattr(args, 'niveau', None):
        cadrage["niveau"] = args.niveau

    return cadrage


def list_templates() -> None:
    """
    Affiche les templates de cadrage disponibles.

    PRD-020: Commande --list-templates.
    """
    print("\n📋 Templates de cadrage disponibles:\n")
    for name, template in CADRAGE_TEMPLATES.items():
        print(f"  {name}:")
        for key, value in template.items():
            print(f"    {key}: {value}")
        print()


# ══════════════════════════════════════════════════════════════════════════════
# PREVIEW ET CONFIRMATION (PRD-021)
# ══════════════════════════════════════════════════════════════════════════════

def display_draft(session: "SyntheseSession") -> None:
    """
    Affiche le brouillon de la synthèse pour validation.

    PRD-021: Preview avant confirmation.
    """
    print("\n" + "═" * 64)
    print("                    📝 BROUILLON")
    print("═" * 64)
    print(f"\nSession: {session.session_id}")
    print(f"Mode: {session.mode} | Modèles: {', '.join(session.models_used)}")
    print(f"Convergence: {session.convergence:.0%} | Durée: {session.duration_total:.1f}s")

    # Cadrage
    print("\n┌─ Cadrage ─────────────────────────────────────────────────┐")
    for key, value in session.cadrage.items():
        if not key.startswith("_"):
            print(f"│  {key}: {value}")
    print("└────────────────────────────────────────────────────────────┘")

    # Synthèse
    print("\n┌─ Synthèse ────────────────────────────────────────────────┐")
    for line in session.synthese_finale.split("\n"):
        # Tronquer les lignes trop longues
        if len(line) > 58:
            line = line[:55] + "..."
        print(f"│  {line}")
    print("└────────────────────────────────────────────────────────────┘")

    print("\n" + "═" * 64)


def prompt_confirmation() -> str:
    """
    Affiche le menu de confirmation et retourne le choix.

    PRD-021: Boucle de validation interactive.

    Returns:
        'export', 'regenerate', 'modify', ou 'cancel'
    """
    print("\n📋 Que souhaitez-vous faire ?")
    print()
    print("  [E] Exporter      - Sauvegarder la synthèse")
    print("  [R] Régénérer     - Relancer la synthèse")
    print("  [M] Modifier      - Ajuster le cadrage")
    print("  [A] Annuler       - Quitter sans sauvegarder")
    print()

    while True:
        try:
            choice = input("Votre choix [E/R/M/A] : ").strip().upper()
        except EOFError:
            return 'export'  # Non-interactif, exporter par défaut

        if choice in ['E', 'EXPORT', 'EXPORTER']:
            return 'export'
        elif choice in ['R', 'REGENERATE', 'REGENERER', 'RÉGÉNÉRER']:
            return 'regenerate'
        elif choice in ['M', 'MODIFY', 'MODIFIER']:
            return 'modify'
        elif choice in ['A', 'ANNULER', 'CANCEL', 'Q', 'QUIT']:
            return 'cancel'
        else:
            print("❌ Choix invalide. Entrez E, R, M ou A.")


def prompt_modify_cadrage(current_cadrage: Dict[str, str]) -> Dict[str, str]:
    """
    Permet de modifier le cadrage existant.

    PRD-021: Ajustement interactif du cadrage.

    Args:
        current_cadrage: Cadrage actuel

    Returns:
        Cadrage modifié
    """
    print("\n📝 Modification du cadrage (Entrée = garder la valeur actuelle)")
    print()

    new_cadrage = {}
    fields = ["destinataire", "finalite", "longueur", "ton", "niveau"]

    for field in fields:
        current = current_cadrage.get(field, "")
        try:
            new_value = input(f"  {field.capitalize()} [{current}] : ").strip()
        except EOFError:
            new_value = ""
        new_cadrage[field] = new_value if new_value else current

    # Conserver les métadonnées
    if "_template" in current_cadrage:
        new_cadrage["_template"] = current_cadrage["_template"]

    return new_cadrage


def generate_frontmatter(session: "SyntheseSession") -> str:
    """
    Genere le frontmatter YAML depuis la session.

    PRD-022: Metadonnees YAML pour ecosystemes statiques.

    Args:
        session: Session de synthese complete

    Returns:
        Frontmatter YAML avec delimiteurs ---
    """
    import yaml

    source_name = "unknown"
    if hasattr(session, 'sources_info') and session.sources_info:
        sources = session.sources_info.get("sources", [])
        if len(sources) == 1:
            source_name = sources[0].get("name", "unknown")
        elif len(sources) > 1:
            source_name = [s.get("name", "unknown") for s in sources]

    metadata = {
        "session_id": session.session_id,
        "source": source_name,
        "date": session.timestamp_start[:10] if session.timestamp_start else None,
        "modeles": session.models_used if hasattr(session, 'models_used') else [],
        "mode": session.mode if hasattr(session, 'mode') else "unknown",
        "convergence": round(session.convergence, 2) if hasattr(session, 'convergence') and session.convergence else 0,
        "duree": round(session.duration_total, 1) if hasattr(session, 'duration_total') and session.duration_total else 0,
    }

    if hasattr(session, 'cadrage') and session.cadrage:
        metadata["cadrage"] = {k: v for k, v in session.cadrage.items() if not k.startswith("_")}
        if session.cadrage.get("_template"):
            metadata["template"] = session.cadrage["_template"]

    yaml_content = yaml.dump(metadata, default_flow_style=False, allow_unicode=True, sort_keys=False)
    return f"---\n{yaml_content}---\n\n"


# ══════════════════════════════════════════════════════════════════════════════
# CACHE DES RÉPONSES (PRD-024)
# ══════════════════════════════════════════════════════════════════════════════

class ResponseCache:
    """
    Cache persistant pour les réponses LLM.

    PRD-024: Éviter les appels redondants en cas d'interruption.
    """

    def __init__(
        self,
        cache_dir: Optional[Path] = None,
        default_ttl: int = 86400,
        enabled: bool = True
    ):
        self.cache_dir = cache_dir or Path.home() / ".synthese_cache"
        self.default_ttl = default_ttl
        self.enabled = enabled
        # PRD-029 Story 29.1: Métriques de cache
        self.hits = 0
        self.misses = 0
        self.total_latency_saved_ms = 0.0
        self._avg_call_duration_ms = 30000.0  # Estimation: 30s par appel
        if self.enabled:
            self.cache_dir.mkdir(parents=True, exist_ok=True)

    def _compute_key(self, model: str, prompt: str, role: str, cadrage: Dict) -> str:
        """Calcule la clé de cache unique."""
        import hashlib
        normalized = {k: v for k, v in sorted(cadrage.items()) if not k.startswith("_")}
        key_data = json.dumps({
            "model": model,
            "prompt": prompt[:500],
            "role": role,
            "cadrage": normalized
        }, sort_keys=True)
        return hashlib.sha256(key_data.encode()).hexdigest()[:16]

    def get(self, model: str, prompt: str, role: str, cadrage: Dict) -> Optional[str]:
        """Récupère une réponse du cache."""
        if not self.enabled:
            return None
        key = self._compute_key(model, prompt, role, cadrage)
        cache_path = self.cache_dir / f"{key}.json"
        if not cache_path.exists():
            self.misses += 1  # PRD-029: Track miss
            return None
        try:
            with open(cache_path, encoding="utf-8") as f:
                data = json.load(f)
            if data.get("ttl", 0) > 0 and time.time() - data.get("timestamp", 0) > data["ttl"]:
                cache_path.unlink()
                self.misses += 1  # PRD-029: Track miss (expired)
                return None
            # PRD-029: Track hit and estimate latency saved
            self.hits += 1
            self.total_latency_saved_ms += self._avg_call_duration_ms
            return data.get("content")
        except Exception:
            self.misses += 1  # PRD-029: Track miss (error)
            return None

    def set(
        self,
        model: str,
        prompt: str,
        role: str,
        cadrage: Dict,
        content: str,
        ttl: Optional[int] = None
    ) -> None:
        """Stocke une réponse dans le cache."""
        if not self.enabled:
            return
        key = self._compute_key(model, prompt, role, cadrage)
        cache_path = self.cache_dir / f"{key}.json"
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({
                "content": content,
                "model": model,
                "timestamp": time.time(),
                "ttl": ttl or self.default_ttl
            }, f, ensure_ascii=False)

    def clear(self) -> int:
        """Vide le cache. Retourne le nombre d'entrées supprimées."""
        count = 0
        if self.cache_dir.exists():
            for f in self.cache_dir.glob("*.json"):
                f.unlink()
                count += 1
        return count

    def stats(self) -> Dict[str, Any]:
        """Retourne les statistiques du cache (PRD-029 Story 29.1)."""
        if not self.cache_dir.exists():
            entries_count = 0
            total_size = 0
        else:
            entries = list(self.cache_dir.glob("*.json"))
            entries_count = len(entries)
            total_size = sum(f.stat().st_size for f in entries)

        total_requests = self.hits + self.misses
        hit_rate = self.hits / total_requests if total_requests > 0 else 0.0

        return {
            "entries": entries_count,
            "size_mb": round(total_size / 1024 / 1024, 2),
            "cache_dir": str(self.cache_dir),
            # PRD-029: Nouvelles métriques
            "hits": self.hits,
            "misses": self.misses,
            "hit_rate": round(hit_rate, 3),
            "latency_saved_ms": round(self.total_latency_saved_ms, 0),
            "latency_saved_s": round(self.total_latency_saved_ms / 1000, 1)
        }


# ══════════════════════════════════════════════════════════════════════════════
# EXTRACTEURS DE CONTENU (PRD-018)
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf(file_path: str) -> str:
    """
    Extrait le texte d'un fichier PDF.

    PRD-018: Support des fichiers PDF via pdfplumber.

    Args:
        file_path: Chemin vers le fichier PDF

    Returns:
        Texte extrait du PDF

    Raises:
        RuntimeError: Si pdfplumber n'est pas installé
        FileNotFoundError: Si le fichier n'existe pas
    """
    if not PDF_AVAILABLE:
        raise RuntimeError(
            "pdfplumber non installé. Installez avec: pip install pdfplumber"
        )

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Fichier PDF non trouvé: {file_path}")

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n\n".join(text_parts)


def extract_docx(file_path: str) -> str:
    """
    Extrait le texte d'un fichier DOCX.

    PRD-018: Support des fichiers Word via python-docx.

    Args:
        file_path: Chemin vers le fichier DOCX

    Returns:
        Texte extrait du document

    Raises:
        RuntimeError: Si python-docx n'est pas installé
        FileNotFoundError: Si le fichier n'existe pas
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx non installé. Installez avec: pip install python-docx"
        )

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Fichier DOCX non trouvé: {file_path}")

    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    return "\n\n".join(paragraphs)


def extract_url(url: str, timeout: int = 30) -> Tuple[str, str]:
    """
    Extrait le texte d'une page web.

    PRD-018: Support des URLs via requests + BeautifulSoup.

    Args:
        url: URL de la page web
        timeout: Timeout en secondes

    Returns:
        Tuple (texte extrait, titre de la page)

    Raises:
        RuntimeError: Si requests/beautifulsoup4 non installés
        requests.RequestException: Si erreur réseau
    """
    if not URL_AVAILABLE:
        raise RuntimeError(
            "requests/beautifulsoup4 non installés. "
            "Installez avec: pip install requests beautifulsoup4"
        )

    headers = {
        'User-Agent': 'Mozilla/5.0 (compatible; SyntheseBot/1.0)'
    }

    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()

    soup = BeautifulSoup(response.content, 'html.parser')

    # Extraire le titre
    title = soup.title.string if soup.title else url

    # Supprimer scripts, styles, etc.
    for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
        element.decompose()

    # Extraire le texte principal
    # Chercher d'abord dans article, main, ou body
    main_content = soup.find('article') or soup.find('main') or soup.body

    if main_content:
        text = main_content.get_text(separator='\n', strip=True)
    else:
        text = soup.get_text(separator='\n', strip=True)

    # Nettoyer les lignes vides multiples
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    text = '\n'.join(lines)

    return text, title


def load_sources(
    files: Optional[List[str]] = None,
    urls: Optional[List[str]] = None,
    text: Optional[str] = None
) -> List[Source]:
    """
    Charge toutes les sources de contenu.

    PRD-018: Fonction unifiée pour charger fichiers, URLs et texte.

    Args:
        files: Liste de chemins de fichiers (txt, pdf, docx)
        urls: Liste d'URLs
        text: Texte brut optionnel

    Returns:
        Liste de Sources avec leur contenu
    """
    sources = []

    # Charger les fichiers
    if files:
        for file_path in files:
            path = Path(file_path)
            ext = path.suffix.lower()

            try:
                if ext == '.pdf':
                    content = extract_pdf(file_path)
                    source_type = "pdf"
                elif ext in ['.docx', '.doc']:
                    content = extract_docx(file_path)
                    source_type = "docx"
                else:
                    # Fichier texte par défaut
                    content = path.read_text(encoding='utf-8')
                    source_type = "file"

                sources.append(Source(
                    name=path.name,
                    content=content,
                    source_type=source_type,
                    metadata={"path": str(path.absolute())}
                ))
            except Exception as e:
                print(f"⚠️  Erreur lecture {file_path}: {e}", file=sys.stderr)

    # Charger les URLs
    if urls:
        for url in urls:
            try:
                content, title = extract_url(url)
                sources.append(Source(
                    name=title,
                    content=content,
                    source_type="url",
                    metadata={"url": url}
                ))
            except Exception as e:
                print(f"⚠️  Erreur URL {url}: {e}", file=sys.stderr)

    # Ajouter le texte brut
    if text:
        sources.append(Source(
            name="Texte direct",
            content=text,
            source_type="text",
            metadata={}
        ))

    return sources


def format_sources_for_prompt(
    sources: List[Source],
    mode: str = "merge"
) -> str:
    """
    Formate les sources pour inclusion dans le prompt.

    PRD-018: Prépare le texte combiné ou comparé selon le mode.

    Args:
        sources: Liste des sources chargées
        mode: "merge" (fusion) ou "compare" (comparaison)

    Returns:
        Texte formaté pour le prompt
    """
    if not sources:
        return ""

    if len(sources) == 1:
        return sources[0].content

    if mode == "compare":
        # Mode comparaison: identifier clairement chaque source
        parts = []
        for i, source in enumerate(sources, 1):
            header = f"═══ SOURCE {i}: {source.name} ({source.source_type}) ═══"
            parts.append(f"{header}\n\n{source.content}")

        intro = f"[{len(sources)} sources à comparer]\n\n"
        return intro + "\n\n".join(parts)

    else:  # mode == "merge"
        # Mode fusion: combiner avec séparateurs légers
        parts = []
        for source in sources:
            header = f"--- {source.name} ---"
            parts.append(f"{header}\n{source.content}")

        return "\n\n".join(parts)


# Rôles experts pour la synthèse (utilisés si config module non disponible)
EXPERT_ROLES = {
    "extracteur": {
        "name": "L'Extracteur de Substance",
        "focus": "faits, données, thèse centrale",
        "prompt_suffix": "Concentre-toi sur les FAITS et la THÈSE centrale. Identifie les données concrètes, les affirmations vérifiables, et le message principal. Ignore les ornements rhétoriques."
    },
    "critique": {
        "name": "Le Gardien de la Fidélité",
        "focus": "glissements, biais, omissions",
        "prompt_suffix": "Traque les GLISSEMENTS de sens. Vérifie que rien n'est ajouté ni omis. Distingue ce que le texte DIT de ce qu'il SUGGÈRE. Signale tout risque de surinterprétation."
    },
    "architecte": {
        "name": "L'Architecte du Sens",
        "focus": "structure, logique, cohérence",
        "prompt_suffix": "Analyse la STRUCTURE argumentative. Identifie les présupposés, les enchaînements logiques, les tensions. Propose une architecture claire pour la synthèse finale."
    }
}

# PRD-017: Presets de personas par domaine
PERSONA_PRESETS = {
    "tech": [
        {"name": "L'Architecte Logiciel", "focus": "structure, patterns, maintenabilité"},
        {"name": "L'Expert Sécurité", "focus": "vulnérabilités, bonnes pratiques, risques"},
        {"name": "L'Optimisateur", "focus": "performance, scalabilité, efficacité"},
    ],
    "science": [
        {"name": "Le Méthodologue", "focus": "rigueur scientifique, protocole, biais"},
        {"name": "Le Spécialiste Domaine", "focus": "état de l'art, concepts clés, implications"},
        {"name": "Le Statisticien", "focus": "données, significativité, reproductibilité"},
    ],
    "legal": [
        {"name": "Le Juriste", "focus": "conformité légale, clauses, risques juridiques"},
        {"name": "Le Fiscaliste", "focus": "implications fiscales, optimisation, conformité"},
        {"name": "L'Expert Compliance", "focus": "régulations, obligations, sanctions"},
    ],
    "business": [
        {"name": "Le Stratège", "focus": "vision, positionnement, avantage concurrentiel"},
        {"name": "L'Analyste Financier", "focus": "chiffres, ratios, rentabilité"},
        {"name": "L'Expert Opérations", "focus": "faisabilité, ressources, mise en œuvre"},
    ],
    "medical": [
        {"name": "Le Clinicien", "focus": "pratique médicale, diagnostic, traitement"},
        {"name": "L'Épidémiologiste", "focus": "données populationnelles, tendances, facteurs de risque"},
        {"name": "Le Biostatisticien", "focus": "analyses statistiques, significativité, biais"},
    ],
    "default": [
        {"name": "L'Extracteur de Substance", "focus": "faits, données, thèse centrale"},
        {"name": "Le Gardien de la Fidélité", "focus": "glissements, biais, omissions"},
        {"name": "L'Architecte du Sens", "focus": "structure, logique, cohérence"},
    ],
}

# PRD-020: Templates de cadrage prédéfinis
CADRAGE_TEMPLATES = {
    "comite-direction": {
        "destinataire": "comité de direction",
        "finalite": "décision",
        "longueur": "1 page",
        "ton": "formel",
        "niveau": "expert"
    },
    "executive": {
        "destinataire": "direction générale",
        "finalite": "briefing rapide",
        "longueur": "5 lignes",
        "ton": "concis et direct",
        "niveau": "stratégique"
    },
    "technique": {
        "destinataire": "développeurs et ingénieurs",
        "finalite": "documentation technique",
        "longueur": "10-15 lignes",
        "ton": "technique et précis",
        "niveau": "expert"
    },
    "vulgarisation": {
        "destinataire": "grand public",
        "finalite": "information accessible",
        "longueur": "10 lignes",
        "ton": "accessible et clair",
        "niveau": "débutant"
    },
    "academique": {
        "destinataire": "chercheurs et universitaires",
        "finalite": "analyse approfondie",
        "longueur": "2 pages",
        "ton": "académique et rigoureux",
        "niveau": "expert"
    },
    "juridique": {
        "destinataire": "juristes et avocats",
        "finalite": "conformité et risques",
        "longueur": "1 page",
        "ton": "formel et précis",
        "niveau": "expert"
    },
    "formation": {
        "destinataire": "apprenants et étudiants",
        "finalite": "apprentissage",
        "longueur": "15 lignes",
        "ton": "didactique et progressif",
        "niveau": "intermédiaire"
    },
    "presse": {
        "destinataire": "journalistes et médias",
        "finalite": "communication",
        "longueur": "5 lignes",
        "ton": "accrocheur et factuel",
        "niveau": "grand public"
    },
}


class Mode(Enum):
    """Modes de délibération"""
    STANDARD = "standard"      # Processus complet en 5 étapes
    RAPIDE = "rapide"          # Cadrage minimal, synthèse directe
    CRITIQUE = "critique"      # Focus sur les glissements
    PEDAGOGIQUE = "pedagogique"  # Explique chaque étape
    DIRECT = "direct"          # Un seul modèle, sans délibération


@dataclass
class ModelResult:
    """Résultat d'un appel à un modèle"""
    model: str
    role: str
    content: str
    duration: float
    success: bool
    error: Optional[str] = None
    timestamp: str = field(default_factory=lambda: datetime.now().isoformat())
    from_cache: bool = False  # PRD-027 Story 27.3


@dataclass
class SyntheseSession:
    """Session de synthèse complète"""
    session_id: str
    query: str
    source_text: str
    mode: str
    cadrage: Dict[str, str]
    rounds: List[Dict[str, Any]] = field(default_factory=list)
    synthese_finale: str = ""
    convergence: float = 0.0
    duration_total: float = 0.0
    timestamp_start: str = field(default_factory=lambda: datetime.now().isoformat())
    timestamp_end: str = ""
    models_used: List[str] = field(default_factory=list)
    retry_metrics: Dict[str, Any] = field(default_factory=dict)  # PRD-004
    cache_metrics: Dict[str, Any] = field(default_factory=dict)  # PRD-029 Story 29.1
    personas_info: Dict[str, Any] = field(default_factory=dict)  # PRD-017
    sources_info: Dict[str, Any] = field(default_factory=dict)  # PRD-018
    references_info: Dict[str, Any] = field(default_factory=dict)  # PRD-019


class NoModelsAvailableError(RuntimeError):
    """
    Exception levée quand aucun modèle LLM n'est disponible.
    
    PRD-008 Story 8.1: Gestion gracieuse des erreurs de configuration.
    """
    def __init__(self, message: str = None):
        default_message = (
            "Aucun modèle LLM disponible.\n\n"
            "Solutions possibles:\n"
            "  1. Définir ANTHROPIC_API_KEY pour utiliser l'API Anthropic\n"
            "  2. Installer un CLI wrapper (claude, gemini, codex)\n"
            "  3. Configurer un backend dans synthese.config.yaml\n\n"
            "Documentation: https://github.com/your-repo/synthese-council#installation"
        )
        super().__init__(message or default_message)


# ══════════════════════════════════════════════════════════════════════════════
# AFFICHAGE THREAD-SAFE (Story 2.5)
# ══════════════════════════════════════════════════════════════════════════════

class AsyncPrinter:
    """Gère l'affichage concurrent sans mélange de lignes."""

    def __init__(self, file=None):
        self._lock = None  # Créé à la demande pour éviter l'erreur "no event loop"
        self._file = file  # None = stdout, sys.stderr pour logs

    def _get_lock(self):
        """Retourne le lock, le créant si nécessaire."""
        if self._lock is None:
            try:
                self._lock = asyncio.Lock()
            except RuntimeError:
                # Pas d'event loop, créer un dummy lock
                self._lock = asyncio.Lock()
        return self._lock

    async def print(self, message: str, end: str = "\n", flush: bool = True):
        """Affiche un message de manière thread-safe."""
        lock = self._get_lock()
        async with lock:
            print(message, end=end, flush=flush, file=self._file)

    def print_sync(self, message: str, end: str = "\n", flush: bool = True):
        """Version synchrone pour le code non-async."""
        print(message, end=end, flush=flush, file=self._file)

    def set_file(self, file):
        """Change la destination de sortie."""
        self._file = file


# Instance globale pour l'affichage
_printer = AsyncPrinter()

# Destination des logs (stdout par défaut, stderr si sortie structurée)
_log_file = None


def log(message: str, end: str = "\n", flush: bool = True):
    """Affiche un message de log (redirigé vers stderr si sortie structurée)."""
    print(message, end=end, flush=flush, file=_log_file)


# ══════════════════════════════════════════════════════════════════════════════
# UTILITAIRES CLI (synchrones - pour vérification)
# ══════════════════════════════════════════════════════════════════════════════

def check_cli(model: str) -> Tuple[bool, str]:
    """Vérifie si un CLI est disponible (synchrone)"""
    cli_map = {
        "claude": "claude",
        "gemini": "gemini",
        "codex": "codex"
    }
    cli = cli_map.get(model, model)
    
    try:
        result = subprocess.run(
            [cli, "--version"],
            capture_output=True,
            text=True,
            timeout=10
        )
        version = result.stdout.strip() or result.stderr.strip()
        return True, version
    except FileNotFoundError:
        return False, f"{cli} non trouvé"
    except subprocess.TimeoutExpired:
        return False, "timeout"
    except Exception as e:
        return False, str(e)


def check_all_clis() -> Dict[str, Tuple[bool, str]]:
    """
    Vérifie les 3 CLIs disponibles.
    
    PRD-013: Retourne uniquement les CLIs (claude, gemini, codex).
    Les backends API sont gérés séparément par get_available_backends().
    
    Returns:
        Dict[model_name, (available, info_message)]
    """
    results = {}
    
    # Vérifier uniquement les CLIs
    for model in ["claude", "gemini", "codex"]:
        results[model] = check_cli(model)
    
    return results


def get_available_backends() -> List[LLMBackend]:
    """
    Récupère les backends API disponibles.
    
    PRD-008: Support des backends natifs (Anthropic, etc.)
    PRD-013: Utilisé comme fallback si aucun CLI disponible.
    
    Returns:
        Liste des backends disponibles
    """
    if not BACKENDS_MODULE_AVAILABLE:
        return []
    
    return BackendRegistry.get_available()


# ══════════════════════════════════════════════════════════════════════════════
# APPELS MODÈLES - VERSION SYNCHRONE (compatibilité)
# ══════════════════════════════════════════════════════════════════════════════

def call_model(
    model: str,
    prompt: str,
    timeout: int = 300,
    role: str = "default"
) -> ModelResult:
    """Appelle un modèle via son CLI (version synchrone)"""
    
    start_time = time.time()
    
    # PRD-013: Commandes CLI correctes
    # - claude: claude -p "prompt"
    # - gemini: gemini -p "prompt"  
    # - codex: codex exec "prompt" (mode non-interactif)
    #   Note: --skip-git-repo-check requis hors d'un repo git
    cli_map = {
        "claude": ["claude", "-p"],
        "gemini": ["gemini", "-p"],
        "codex": ["codex", "exec", "--skip-git-repo-check"]
    }
    
    cmd = cli_map.get(model)
    if not cmd:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=0,
            success=False,
            error=f"Modèle inconnu: {model}"
        )
    
    try:
        result = subprocess.run(
            cmd + [prompt],
            capture_output=True,
            text=True,
            timeout=timeout
        )
        
        duration = time.time() - start_time
        content = result.stdout.strip()
        
        if result.returncode != 0 and not content:
            return ModelResult(
                model=model,
                role=role,
                content="",
                duration=duration,
                success=False,
                error=result.stderr.strip() or f"Code retour: {result.returncode}"
            )
        
        return ModelResult(
            model=model,
            role=role,
            content=content,
            duration=duration,
            success=True
        )
        
    except subprocess.TimeoutExpired:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=timeout,
            success=False,
            error=f"Timeout après {timeout}s"
        )
    except FileNotFoundError:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=time.time() - start_time,
            success=False,
            error=f"CLI non trouvé: {model}"
        )
    except Exception as e:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=time.time() - start_time,
            success=False,
            error=str(e)
        )


# ══════════════════════════════════════════════════════════════════════════════
# APPELS MODÈLES - VERSION ASYNCHRONE (Story 2.1)
# ══════════════════════════════════════════════════════════════════════════════

# Instance globale de l'executor par défaut (PRD-015)
_default_cli_executor = None


def get_default_cli_executor() -> DefaultCLIExecutor:
    """Retourne l'executor CLI par défaut (singleton)."""
    global _default_cli_executor
    if _default_cli_executor is None:
        _default_cli_executor = DefaultCLIExecutor()
    return _default_cli_executor


async def call_model_async(
    model: str,
    prompt: str,
    timeout: int = 300,
    role: str = "default",
    printer: AsyncPrinter = None,
    cli_executor: CLIExecutor = None
) -> ModelResult:
    """
    Appelle un modèle via son CLI de manière asynchrone.

    Story 2.1 du PRD-002: Version async de call_model avec asyncio.create_subprocess_exec
    PRD-015: Support de l'injection de dépendances via cli_executor

    Args:
        model: Nom du modèle (claude, gemini, codex)
        prompt: Prompt à envoyer
        timeout: Timeout en secondes
        role: Rôle assigné
        printer: Printer async pour l'affichage
        cli_executor: Executor CLI injectable (pour les tests)
    """

    start_time = time.time()
    executor = cli_executor or get_default_cli_executor()

    # PRD-013: Commandes CLI correctes
    # - claude: claude -p "prompt"
    # - gemini: gemini -p "prompt"
    # - codex: codex exec "prompt" (mode non-interactif)
    #   Note: --skip-git-repo-check requis hors d'un repo git
    cli_map = {
        "claude": ["claude", "-p"],
        "gemini": ["gemini", "-p"],
        "codex": ["codex", "exec", "--skip-git-repo-check"]
    }

    cmd = cli_map.get(model)
    if not cmd:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=0,
            success=False,
            error=f"Modèle inconnu: {model}"
        )

    try:
        # Utiliser l'executor injectable (PRD-015)
        stdout, stderr, returncode = await executor.run(cmd + [prompt], timeout)

        duration = time.time() - start_time
        content = stdout.strip()
        stderr_text = stderr.strip()

        if returncode != 0 and not content:
            return ModelResult(
                model=model,
                role=role,
                content="",
                duration=duration,
                success=False,
                error=stderr_text or f"Code retour: {returncode}"
            )

        return ModelResult(
            model=model,
            role=role,
            content=content,
            duration=duration,
            success=True
        )

    except asyncio.TimeoutError:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=timeout,
            success=False,
            error=f"Timeout après {timeout}s"
        )
    except FileNotFoundError:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=time.time() - start_time,
            success=False,
            error=f"CLI non trouvé: {model}"
        )
    except asyncio.CancelledError:
        # Propagate cancellation
        raise
    except Exception as e:
        return ModelResult(
            model=model,
            role=role,
            content="",
            duration=time.time() - start_time,
            success=False,
            error=str(e)
        )


# ══════════════════════════════════════════════════════════════════════════════
# APPELS MODÈLES AVEC RETRY (PRD-004)
# ══════════════════════════════════════════════════════════════════════════════

async def call_model_async_with_retry(
    model: str,
    prompt: str,
    timeout: int = 300,
    role: str = "default",
    printer: AsyncPrinter = None,
    retry_config: Dict[str, Any] = None,
    metrics_collector = None,
    cli_executor: CLIExecutor = None,
    cache: "ResponseCache" = None,
    cadrage: Dict[str, Any] = None
) -> ModelResult:
    """
    Appelle un modèle avec retry automatique sur erreurs transitoires.

    PRD-004: Wrapper de call_model_async avec backoff exponentiel.
    PRD-027 Story 27.3: Intégration cache.

    Args:
        model: Nom du modèle (claude, gemini, codex)
        prompt: Prompt à envoyer
        timeout: Timeout par tentative
        role: Rôle assigné au modèle
        printer: Printer async pour l'affichage
        retry_config: Configuration du retry (depuis config YAML)
        metrics_collector: Collecteur de métriques de retry
        cache: Cache de réponses (PRD-027)
        cadrage: Paramètres de cadrage pour clé de cache

    Returns:
        ModelResult avec succès ou erreur finale
    """
    # PRD-027 Story 27.3: Cache lookup
    if cache and cache.enabled:
        cached_content = cache.get(model, prompt, role, cadrage or {})
        if cached_content:
            if printer:
                await printer.print(f"  ⚡ Cache hit pour {model}")
            return ModelResult(
                model=model,
                content=cached_content,
                success=True,
                duration=0.0,
                role=role,
                from_cache=True
            )

    if not RETRY_MODULE_AVAILABLE:
        # Pas de module retry, appel direct
        result = await call_model_async(model, prompt, timeout, role, printer, cli_executor)
        # Cache store après succès
        if cache and cache.enabled and result.success:
            cache.set(model, prompt, role, cadrage or {}, result.content)
        return result

    # Configuration par défaut
    cfg = {**DEFAULT_RETRY_CONFIG, **(retry_config or {})}

    if not cfg.get("enabled", True):
        # Retry désactivé
        result = await call_model_async(model, prompt, timeout, role, printer, cli_executor)
        if cache and cache.enabled and result.success:
            cache.set(model, prompt, role, cadrage or {}, result.content)
        return result
    
    # Récupérer les paramètres
    max_attempts = cfg.get("max_attempts", 3)
    base_delay = cfg.get("base_delay", 2.0)
    max_delay = cfg.get("max_delay", 30.0)
    exponential_base = cfg.get("exponential_base", 2.0)
    jitter = cfg.get("jitter", 0.1)
    
    # Callback pour l'affichage
    callback = PrintRetryCallback(printer) if printer else None
    
    # Métriques
    collector = metrics_collector or get_retry_metrics()
    
    last_result = None
    
    for attempt in range(1, max_attempts + 1):
        collector.record_attempt(model)

        # Appel au modèle (PRD-015: avec executor injectable)
        result = await call_model_async(model, prompt, timeout, role, printer, cli_executor)
        
        if result.success:
            # Succès
            if callback and attempt > 1:
                await callback.on_success(model, attempt)

            if attempt > 1:
                collector.record_success_after_retry(model)

            # PRD-027 Story 27.3: Cache store après succès
            if cache and cache.enabled:
                cache.set(model, prompt, role, cadrage or {}, result.content)

            return result
        
        # Échec - vérifier si retryable
        last_result = result
        
        if not is_retryable(result.error or ""):
            # Erreur non-retryable (CLI absent, etc.)
            collector.record_failure(model)
            return result
        
        if attempt == max_attempts:
            # Dernière tentative, échec final
            collector.record_failure(model)
            
            if callback:
                await callback.on_failure(model, attempt, classify_error(result.error or ""))
            
            return result
        
        # Calculer le délai
        import random as _random
        delay = base_delay * (exponential_base ** (attempt - 1))
        delay = min(delay, max_delay)
        jitter_range = delay * jitter
        delay += _random.uniform(-jitter_range, jitter_range)
        delay = max(0, delay)
        
        # Logger et notifier
        reason = classify_error(result.error or "")
        # PRD-027 Story 27.2: Événements détaillés
        collector.record_retry(
            model, reason, delay,
            attempt=attempt,
            error_snippet=result.error or ""
        )
        
        if callback:
            await callback.on_retry(model, attempt, max_attempts, delay, reason)
        
        # Attendre avant le retry
        await asyncio.sleep(delay)
    
    return last_result


# ══════════════════════════════════════════════════════════════════════════════
# MOTEUR DE SYNTHÈSE
# ══════════════════════════════════════════════════════════════════════════════

class SyntheseCouncil:
    """Orchestre la synthèse multi-LLM avec appels asynchrones et config dynamique"""

    def __init__(
        self,
        config: Dict = None,
        config_path: str = None,
        cli_executor: CLIExecutor = None,
        backend_factory = None
    ):
        """
        Initialise le council avec configuration.

        PRD-015: Support de l'injection de dépendances pour les tests.

        Args:
            config: Override de configuration (dict)
            config_path: Chemin vers un fichier de config YAML
            cli_executor: Executor CLI injectable (pour les tests)
            backend_factory: Factory de backends injectable (pour les tests)
        """
        # PRD-015: Injection de dépendances
        self.cli_executor = cli_executor
        self.backend_factory = backend_factory
        # Charger la configuration (PRD-003)
        if CONFIG_MODULE_AVAILABLE:
            try:
                self.config = ConfigLoader.load(
                    path=config_path,
                    cli_overrides=config
                )
            except ConfigError as e:
                print(f"⚠ Erreur de configuration: {e}")
                print("  Utilisation de la configuration par défaut.")
                self.config = {**DEFAULT_CONFIG, **(config or {})}
        else:
            self.config = {**DEFAULT_CONFIG, **(config or {})}
        
        # Charger les rôles experts (avec personnalisation)
        if CONFIG_MODULE_AVAILABLE:
            self.expert_roles = get_expert_roles(self.config)
        else:
            self.expert_roles = EXPERT_ROLES
        
        # Configuration retry (PRD-004)
        self.retry_config = self.config.get("retry", DEFAULT_RETRY_CONFIG)
        
        # Métriques de retry (PRD-004)
        if RETRY_MODULE_AVAILABLE:
            self.retry_metrics = RetryMetricsCollector()
        else:
            self.retry_metrics = None
        
        # Mode pédagogique (PRD-005)
        self.pedagogy = None  # Sera initialisé dans run_async selon le mode
        
        self.available_models: List[str] = []
        self.session: Optional[SyntheseSession] = None
        self.printer = _printer  # Utiliser le printer global (PRD-014)
        
        # PRD-008: Backends API natifs
        self.backends: Dict[str, Any] = {}  # model_name -> backend instance

        # PRD-029: Cache de réponses câblé dans flux principal
        cache_config = self.config.get("cache", {})
        self.cache = ResponseCache(
            enabled=cache_config.get("enabled", True),
            default_ttl=cache_config.get("ttl", 3600),
            cache_dir=Path(cache_config.get("directory", ".cache/synthese")) if cache_config.get("directory") else None
        )
        self.cadrage: Dict[str, str] = {}  # Stocké pour clé de cache
        
    def setup(self) -> bool:
        """
        Vérifie et configure les modèles disponibles.
        
        PRD-013: Priorité aux CLIs, API en fallback.
        
        Ordre de priorité:
        1. CLIs (claude, gemini, codex) - méthode principale
        2. Backend API (Anthropic) - fallback si aucun CLI
        """
        log("\n╔══════════════════════════════════════════════════════════════╗")
        log("║           VÉRIFICATION DES MODÈLES                           ║")
        log("╠══════════════════════════════════════════════════════════════╣")

        # 1. Vérifier les CLIs en premier (méthode principale)
        cli_status = check_all_clis()

        for model, (available, info) in cli_status.items():
            status = "✓" if available else "✗"
            log(f"║  {status} {model:10} │ {info[:45]:<45} ║")
            if available:
                self.available_models.append(model)

        # 2. Fallback: Backend API si aucun CLI disponible
        if not self.available_models and BACKENDS_MODULE_AVAILABLE:
            log("╠══════════════════════════════════════════════════════════════╣")
            log("║  Aucun CLI trouvé, recherche de backends API...              ║")
            log("╠══════════════════════════════════════════════════════════════╣")

            api_backends = get_available_backends()
            for backend in api_backends:
                model_name = backend.name.split(":")[0]
                self.backends[model_name] = backend
                self.available_models.append(model_name)
                log(f"║  ✓ {model_name:10} │ {backend.display_name[:45]:<45} ║")

        log("╠══════════════════════════════════════════════════════════════╣")

        if len(self.available_models) == 0:
            log("║  ERREUR: Aucun modèle disponible                            ║")
            log("║                                                              ║")
            log("║  Solutions:                                                  ║")
            log("║  1. Installer les CLIs:                                      ║")
            log("║     npm install -g @anthropic-ai/claude-code                 ║")
            log("║     npm install -g @google/gemini-cli                        ║")
            log("║     npm install -g @openai/codex                             ║")
            log("║  2. Ou définir ANTHROPIC_API_KEY (fallback)                  ║")
            log("╚══════════════════════════════════════════════════════════════╝")
            return False
        elif len(self.available_models) < 3 and not self.backends:
            log(f"║  ATTENTION: {len(self.available_models)}/3 CLIs disponibles                          ║")
            if not self.config["fallback_to_available"]:
                log("╚══════════════════════════════════════════════════════════════╝")
                return False

        # Afficher le type de backend utilisé
        if self.backends:
            log(f"║  STATUS: Prêt avec {len(self.available_models)} modèle(s) (API fallback) ✓         ║")
        else:
            log(f"║  STATUS: Prêt avec {len(self.available_models)} modèle(s) (CLI) ✓                  ║")
        log("╚══════════════════════════════════════════════════════════════╝")
        return True
    
    def _generate_session_id(self) -> str:
        """Génère un ID de session unique"""
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        hash_suffix = hashlib.md5(str(time.time()).encode()).hexdigest()[:6]
        return f"synthese-{timestamp}-{hash_suffix}"
    
    def _assign_roles(self) -> Dict[str, str]:
        """Assigne les rôles aux modèles disponibles"""
        roles = list(self.expert_roles.keys())
        assignments = {}

        for i, model in enumerate(self.available_models):
            role = roles[i % len(roles)]
            assignments[model] = role

        return assignments

    # ══════════════════════════════════════════════════════════════════════════
    # PERSONAS DYNAMIQUES (PRD-017)
    # ══════════════════════════════════════════════════════════════════════════

    def _preset_to_roles(self, preset: List[Dict]) -> Dict[str, Dict]:
        """Convertit un preset de personas en structure expert_roles (PRD-017)."""
        roles = {}
        for i, persona in enumerate(preset):
            role_key = f"persona_{i}"
            roles[role_key] = {
                "name": persona["name"],
                "focus": persona["focus"],
                "prompt_suffix": f"En tant que {persona['name']}, concentre-toi sur: {persona['focus']}. Apporte ta perspective d'expert sur ces aspects spécifiques."
            }
        return roles

    def _custom_to_roles(self, names: List[str]) -> Dict[str, Dict]:
        """Convertit une liste de noms personnalisés en structure expert_roles (PRD-017)."""
        roles = {}
        for i, name in enumerate(names[:3]):  # Max 3 personas
            role_key = f"persona_{i}"
            clean_name = name.strip().capitalize()
            roles[role_key] = {
                "name": f"L'Expert {clean_name}",
                "focus": f"perspective {clean_name}",
                "prompt_suffix": f"En tant qu'expert {clean_name}, analyse ce document avec ta perspective spécialisée. Apporte ton regard critique et tes connaissances du domaine."
            }
        return roles

    async def generate_personas(self, source_text: str, model: str = None) -> List[Dict[str, str]]:
        """
        Génère des personas experts adaptés au document (PRD-017).

        Args:
            source_text: Texte à analyser
            model: Modèle à utiliser (défaut: premier disponible)

        Returns:
            Liste de 3 personas avec name et focus
        """
        model = model or self.available_models[0]

        prompt = f"""Analyse ce document et génère 3 experts complémentaires pour le synthétiser.

DOCUMENT (extrait):
---
{source_text[:2000]}
---

CONSIGNES:
1. Identifie le domaine principal du document
2. Génère 3 personas d'experts avec des angles complémentaires
3. Chaque persona doit avoir un nom évocateur et un focus précis

FORMAT DE RÉPONSE (JSON strict):
{{
  "domaine": "nom du domaine identifié",
  "personas": [
    {{"name": "Le/La [Titre Expert 1]", "focus": "aspect 1, aspect 2, aspect 3"}},
    {{"name": "Le/La [Titre Expert 2]", "focus": "aspect 1, aspect 2, aspect 3"}},
    {{"name": "Le/La [Titre Expert 3]", "focus": "aspect 1, aspect 2, aspect 3"}}
  ]
}}

Réponds UNIQUEMENT avec le JSON, sans texte avant ou après."""

        result = await self._call_model_smart(model, prompt, timeout=60, role="persona_generator")

        if not result.success:
            log(f"⚠ Génération personas échouée, utilisation des défauts")
            return PERSONA_PRESETS["default"]

        try:
            # Parser le JSON
            import re
            json_match = re.search(r'\{[\s\S]*\}', result.content)
            if json_match:
                data = json.loads(json_match.group())
                personas = data.get("personas", [])
                domaine = data.get("domaine", "inconnu")
                if len(personas) >= 3:
                    log(f"✓ Domaine détecté: {domaine}")
                    # Stocker le domaine détecté pour le trail
                    self._detected_domain = domaine
                    return personas[:3]
        except (json.JSONDecodeError, KeyError) as e:
            log(f"⚠ Erreur parsing personas: {e}")

        return PERSONA_PRESETS["default"]

    async def resolve_personas(self, personas_arg: str, source_text: str) -> Tuple[Dict[str, Dict], str, List[Dict]]:
        """
        Résout les personas selon l'argument fourni (PRD-017).

        Args:
            personas_arg: 'auto', nom de preset, ou liste personnalisée
            source_text: Texte source pour l'auto-détection

        Returns:
            Tuple (expert_roles, mode_label, personas_list)
        """
        self._detected_domain = None

        if personas_arg == "auto":
            log("\n─── GÉNÉRATION DES PERSONAS ───")
            log("├─ Analyse du document...")
            personas_list = await self.generate_personas(source_text)
            roles = self._preset_to_roles(personas_list)
            return roles, "auto", personas_list

        elif personas_arg in PERSONA_PRESETS:
            personas_list = PERSONA_PRESETS[personas_arg]
            roles = self._preset_to_roles(personas_list)
            log(f"✓ Personas preset: {personas_arg}")
            return roles, personas_arg, personas_list

        else:
            # Liste personnalisée "expert1,expert2,expert3"
            custom_names = [p.strip() for p in personas_arg.split(",")]
            roles = self._custom_to_roles(custom_names)
            personas_list = [{"name": f"L'Expert {n.strip().capitalize()}", "focus": f"perspective {n.strip()}"} for n in custom_names[:3]]
            log(f"✓ Personas personnalisés: {', '.join(custom_names[:3])}")
            return roles, "custom", personas_list

    async def _call_model_smart(
        self,
        model: str,
        prompt: str,
        timeout: int = 300,
        role: str = "default"
    ) -> ModelResult:
        """
        Appelle un modèle via le backend approprié (API ou CLI).
        
        PRD-008 Story 8.2: Utilise le backend API si disponible,
        sinon fallback sur CLI avec retry.
        
        Args:
            model: Nom du modèle
            prompt: Prompt à envoyer
            timeout: Timeout en secondes
            role: Rôle assigné
            
        Returns:
            ModelResult avec le résultat
        """
        start_time = time.time()
        
        # Vérifier si un backend API est disponible pour ce modèle
        if model in self.backends:
            backend = self.backends[model]
            
            try:
                # Appel via backend API
                response = await backend.call(prompt, timeout=timeout)
                
                return ModelResult(
                    model=model,
                    role=role,
                    content=response.content,
                    duration=response.duration,
                    success=response.success,
                    error=response.error
                )
                
            except Exception as e:
                duration = time.time() - start_time
                return ModelResult(
                    model=model,
                    role=role,
                    content="",
                    duration=duration,
                    success=False,
                    error=f"Erreur backend {model}: {e}"
                )
        
        # Fallback sur CLI avec retry (PRD-015: avec executor injectable)
        # PRD-029: Cache câblé dans flux principal
        return await call_model_async_with_retry(
            model,
            prompt,
            timeout,
            role,
            printer=self.printer,
            retry_config=self.retry_config,
            metrics_collector=self.retry_metrics,
            cli_executor=self.cli_executor,
            cache=self.cache,
            cadrage=self.cadrage
        )
    
    def _build_prompt(
        self,
        stage: str,
        source_text: str,
        cadrage: Dict[str, str],
        role: str,
        previous_outputs: List[Dict] = None,
        with_refs: bool = False  # PRD-019
    ) -> str:
        """Construit le prompt pour une étape donnée"""

        role_info = self.expert_roles.get(role, {})
        role_name = role_info.get("name", role)
        role_suffix = role_info.get("prompt_suffix", "")
        
        base_prompt = f"""Tu es {role_name}, expert en analyse de texte.

CONTEXTE DE LA SYNTHÈSE:
- Destinataire: {cadrage.get('destinataire', 'non précisé')}
- Finalité: {cadrage.get('finalite', 'non précisé')}
- Longueur: {cadrage.get('longueur', 'non précisé')}
- Ton: {cadrage.get('ton', 'non précisé')}
- Niveau: {cadrage.get('niveau', 'non précisé')}

TEXTE SOURCE:
---
{source_text}
---

"""
        
        if stage == "extraction":
            base_prompt += f"""TÂCHE: Extraction initiale

{role_suffix}

Produis une analyse structurée selon ton expertise. Format attendu:
1. THÈSE CENTRALE (1-2 phrases)
2. FAITS CLÉS (liste)
3. MOTS-CADRES (termes importants)
4. LOGIQUE ARGUMENTATIVE (structure)
5. POINTS DE VIGILANCE (risques de glissement)
"""
        
        elif stage == "critique":
            prev_content = "\n\n".join([
                f"=== {p['model']} ({p['role']}) ===\n{p['content']}"
                for p in (previous_outputs or [])
            ])
            
            base_prompt += f"""TÂCHE: Critique croisée

Les autres experts ont produit les analyses suivantes:

{prev_content}

{role_suffix}

Évalue les analyses précédentes:
1. CONVERGENCES (points d'accord)
2. DIVERGENCES (désaccords à résoudre)
3. GLISSEMENTS DÉTECTÉS (erreurs potentielles)
4. ENRICHISSEMENTS (ce qui manque)
5. VERDICT (fiabilité globale 0-10)
"""
        
        elif stage == "synthese":
            prev_content = "\n\n".join([
                f"=== {p['model']} ({p['role']}) ===\n{p['content']}"
                for p in (previous_outputs or [])
            ])

            # PRD-019: Instructions de référencement
            refs_instructions = ""
            if with_refs:
                refs_instructions = """
RÈGLES DE RÉFÉRENCEMENT (OBLIGATOIRE):
- Chaque fait ou chiffre DOIT avoir une référence [§N]
- N = numéro du paragraphe source (voir [§1], [§2], etc. dans le texte)
- Format: "Le CA a augmenté de 15% [§3]"
- Ne référence que ce qui est DANS le texte source
- Pas de référence pour conclusions/interprétations personnelles
- Multi-sources: utilise [S:nom§N] si plusieurs sources
"""

            base_prompt += f"""TÂCHE: Synthèse finale

Voici les analyses et critiques des experts:

{prev_content}

CONSIGNES:
- Longueur: {cadrage.get('longueur', '10-15 lignes')}
- Ton: {cadrage.get('ton', 'sobre')}
- Ne retenir QUE ce que le texte affirme
- Signaler les interprétations si nécessaire
- Mentionner les points de dissensus importants
{refs_instructions}
Produis la SYNTHÈSE FINALE.
"""
        
        return base_prompt
    
    def _calculate_convergence(self, results: List[ModelResult]) -> int:
        """
        Calcule le score de convergence entre les modèles.

        PRD-007: Utilise l'algorithme hybride si disponible.
        PRD-027 Story 27.1: Support sémantique optionnel.
        Combine Jaccard n-grams + cosine similarity pour une mesure robuste.

        Returns:
            Score de convergence en pourcentage (0-100)
        """
        if len(results) < 2:
            return 100

        # Extraire les contenus
        contents = [r.content for r in results if r.content]

        if len(contents) < 2:
            return 100

        # Utiliser le module hybride si disponible (PRD-007, PRD-027)
        if CONVERGENCE_MODULE_AVAILABLE:
            weights = self.config.get("convergence_weights", None)
            result = calc_convergence_hybrid(contents, weights, self.config)
            # PRD-027: result est maintenant un dict
            if isinstance(result, dict):
                # Stocker les détails pour les métriques
                self._convergence_details = result
                return result.get("score", 0)
            # Compatibilité avec ancien format (float)
            return int(result * 100)
        
        # Fallback: algorithme Jaccard simple
        all_words = []
        for content in contents:
            words = set(re.findall(r'\b\w{4,}\b', content.lower()))
            all_words.append(words)
        
        if not all_words:
            return 0.5
        
        # Intersection / Union (Jaccard simplifié)
        common = all_words[0]
        total = all_words[0]
        for words in all_words[1:]:
            common = common & words
            total = total | words
        
        if not total:
            return 0.5
        
        return len(common) / len(total)
    
    # ══════════════════════════════════════════════════════════════════════════
    # EXECUTION PARALLÈLE DES ROUNDS (Story 2.2)
    # ══════════════════════════════════════════════════════════════════════════
    
    async def _run_round_async(
        self,
        stage: str,
        source_text: str,
        cadrage: Dict[str, str],
        role_assignments: Dict[str, str],
        previous_outputs: List[Dict] = None
    ) -> List[Dict[str, Any]]:
        """
        Exécute un round avec tous les modèles EN PARALLÈLE.
        
        Story 2.2 du PRD-002: Utilise asyncio.gather pour paralléliser les appels.
        PRD-004: Utilise call_model_async_with_retry pour la résilience.
        """
        
        async def call_single_model(model: str) -> Dict[str, Any]:
            """Appelle un modèle et retourne le résultat formaté."""
            role = role_assignments[model]
            role_name = self.expert_roles[role]["name"]
            
            # Affichage du démarrage (thread-safe)
            await self.printer.print(f"├─ {model} ({role_name})...", end=" ", flush=True)
            
            # Construire le prompt
            if stage == "critique":
                # Pour la critique, exclure sa propre sortie
                others = [r for r in (previous_outputs or []) if r["model"] != model and r.get("success", True)]
                prompt = self._build_prompt(stage, source_text, cadrage, role, others)
            else:
                prompt = self._build_prompt(stage, source_text, cadrage, role, previous_outputs)
            
            # Appel asynchrone avec backend smart (PRD-008) + retry (PRD-004)
            result = await self._call_model_smart(
                model, 
                prompt, 
                self.config["timeout"], 
                role
            )
            
            # Affichage du résultat (thread-safe)
            if result.success:
                await self.printer.print(f"✓ ({result.duration:.1f}s)")
            else:
                await self.printer.print(f"✗ ({result.error})")
            
            return {
                "model": model,
                "role": role,
                "content": result.content,
                "duration": result.duration,
                "success": result.success,
                "error": result.error
            }
        
        # Exécuter tous les modèles en parallèle
        tasks = [call_single_model(model) for model in self.available_models]
        
        # asyncio.gather avec return_exceptions=True pour ne pas interrompre si un échoue
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Traiter les résultats
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                # Un modèle a levé une exception
                model = self.available_models[i]
                processed_results.append({
                    "model": model,
                    "role": role_assignments[model],
                    "content": "",
                    "duration": 0,
                    "success": False,
                    "error": str(result)
                })
            else:
                processed_results.append(result)
        
        return processed_results
    
    # ══════════════════════════════════════════════════════════════════════════
    # MÉTHODE RUN ASYNCHRONE (Story 2.3)
    # ══════════════════════════════════════════════════════════════════════════
    
    async def run_async(
        self,
        source_text: str,
        cadrage: Dict[str, str],
        mode: Mode = Mode.STANDARD,
        interactive_pedagogy: bool = False,
        personas: str = None,
        with_refs: bool = False  # PRD-019
    ) -> SyntheseSession:
        """
        Exécute le processus complet de synthèse de manière ASYNCHRONE.

        Story 2.3 du PRD-002: Version async de run() avec appels parallèles.
        PRD-005: Mode pédagogique intégré.
        PRD-017: Personas dynamiques.
        PRD-019: Références aux sources.

        Args:
            source_text: Texte à synthétiser
            cadrage: Paramètres de cadrage
            mode: Mode de synthèse (STANDARD, RAPIDE, CRITIQUE, PEDAGOGIQUE)
            interactive_pedagogy: Pause après chaque explication pédagogique
            personas: 'auto', preset (tech/science/legal/business/medical), ou liste personnalisée
            with_refs: Inclure des références [§N] aux passages sources
        """

        start_time = time.time()
        personas_info = None  # Pour le trail
        references_info = None  # PRD-019
        
        # PRD-008 Story 8.1: Vérifier qu'au moins un modèle est disponible
        if not self.available_models:
            raise NoModelsAvailableError()
        
        # PRD-008 Story 8.4: Sanitiser les entrées
        if SANITIZE_MODULE_AVAILABLE:
            clean_text, clean_cadrage, _, warnings = validate_and_sanitize(
                source_text, cadrage, strict=False
            )
            if warnings:
                for w in warnings:
                    print(f"⚠ {w}")
            source_text = clean_text
            cadrage = clean_cadrage

        # PRD-029: Stocker cadrage pour clé de cache
        self.cadrage = cadrage

        # PRD-017: Résoudre les personas dynamiques
        if personas:
            new_roles, personas_mode, personas_list = await self.resolve_personas(personas, source_text)
            self.expert_roles = new_roles
            personas_info = {
                "mode": personas_mode,
                "domain": getattr(self, '_detected_domain', None),
                "personas": personas_list
            }
            # Afficher les personas générés
            log("└─ Personas générés:")
            for p in personas_list:
                log(f"   • {p['name']} ({p['focus']})")

        # Initialiser le mode pédagogique (PRD-005)
        if PEDAGOGY_MODULE_AVAILABLE:
            self.pedagogy = create_pedagogy(mode.value, interactive_pedagogy)
        else:
            self.pedagogy = None
        
        # Afficher l'introduction pédagogique
        if self.pedagogy:
            self.pedagogy.explain(PedagogicStep.INTRO)
        
        # Initialisation de la session
        self.session = SyntheseSession(
            session_id=self._generate_session_id(),
            query="synthèse",
            source_text=source_text[:500] + "..." if len(source_text) > 500 else source_text,
            mode=mode.value,
            cadrage=cadrage,
            models_used=self.available_models.copy()
        )
        
        log(f"\nSession: {self.session.session_id}")
        log(f"Mode: {mode.value} | Modèles: {', '.join(self.available_models)}")
        log(f"⚡ Exécution parallèle activée")
        
        role_assignments = self._assign_roles()
        
        # ═══════════════════════════════════════════════════════════════════
        # ROUND 1: EXTRACTION (parallèle)
        # ═══════════════════════════════════════════════════════════════════
        
        # Explication pédagogique de l'extraction (PRD-005)
        if self.pedagogy:
            self.pedagogy.explain(PedagogicStep.EXTRACTION)
            # Expliquer chaque expert
            for model in self.available_models:
                role = role_assignments[model]
                self.pedagogy.explain_expert(role, model)
        
        log("\n─── ROUND 1: EXTRACTION (parallèle) ───")
        round1_start = time.time()
        
        round1_results = await self._run_round_async(
            "extraction",
            source_text,
            cadrage,
            role_assignments
        )
        
        round1_duration = time.time() - round1_start
        log(f"└─ Round 1 terminé en {round1_duration:.1f}s")
        
        self.session.rounds.append({
            "stage": "extraction",
            "results": round1_results,
            "duration": round1_duration
        })
        
        # ═══════════════════════════════════════════════════════════════════
        # ROUND 2: CRITIQUE CROISÉE (parallèle, si applicable)
        # ═══════════════════════════════════════════════════════════════════
        
        successful_round1 = [r for r in round1_results if r["success"]]
        
        if mode != Mode.RAPIDE and len(successful_round1) >= 2:
            # Explication pédagogique de la critique (PRD-005)
            if self.pedagogy:
                self.pedagogy.explain(PedagogicStep.CRITIQUE)
                self.pedagogy.explain(PedagogicStep.GLISSEMENTS)
            
            log("\n─── ROUND 2: CRITIQUE CROISÉE (parallèle) ───")
            round2_start = time.time()
            
            round2_results = await self._run_round_async(
                "critique",
                source_text,
                cadrage,
                role_assignments,
                round1_results
            )
            
            round2_duration = time.time() - round2_start
            log(f"└─ Round 2 terminé en {round2_duration:.1f}s")
            
            self.session.rounds.append({
                "stage": "critique",
                "results": round2_results,
                "duration": round2_duration
            })
            
            # Calcul de convergence
            successful_results = [
                ModelResult(
                    model=r["model"],
                    role=r["role"],
                    content=r["content"],
                    duration=r["duration"],
                    success=r["success"]
                )
                for r in round2_results if r["success"]
            ]
            self.session.convergence = self._calculate_convergence(successful_results)
            log(f"    Convergence: {self.session.convergence:.2f}")
            
            # Explication pédagogique de la convergence (PRD-005)
            if self.pedagogy:
                self.pedagogy.explain(PedagogicStep.CONVERGENCE)
                self.pedagogy.show_convergence(self.session.convergence)
        
        # ═══════════════════════════════════════════════════════════════════
        # ROUND FINAL: SYNTHÈSE
        # ═══════════════════════════════════════════════════════════════════
        
        # Explication pédagogique de la synthèse (PRD-005)
        if self.pedagogy:
            self.pedagogy.explain(PedagogicStep.SYNTHESE)
        
        log("\n─── SYNTHÈSE FINALE ───")

        # Rassembler tous les outputs réussis
        all_outputs = []
        for round_data in self.session.rounds:
            for r in round_data["results"]:
                if r["success"]:
                    all_outputs.append(r)

        # Utiliser le premier modèle disponible pour la synthèse
        synthesizer = self.available_models[0]
        log(f"├─ {synthesizer} (synthétiseur)...", end=" ")

        # PRD-019: Indexer le texte source si références demandées
        indexed_source = None
        source_for_prompt = source_text
        if with_refs:
            indexed_source = index_source(source_text)
            source_for_prompt = format_indexed_for_prompt(indexed_source)

        prompt = self._build_prompt(
            "synthese", source_for_prompt, cadrage, "architecte", all_outputs, with_refs=with_refs
        )

        # Appel avec backend smart (PRD-008) + retry (PRD-004)
        result = await self._call_model_smart(
            synthesizer,
            prompt,
            self.config["timeout"],
            "synthese"
        )

        if result.success:
            log(f"✓ ({result.duration:.1f}s)")
            self.session.synthese_finale = result.content

            # PRD-019: Valider les références
            if with_refs and indexed_source:
                valid, errors, valid_refs = validate_references(
                    result.content, indexed_source
                )
                if errors:
                    log(f"⚠️  Références invalides: {', '.join(errors)}")

                # Générer les extraits
                indexed_dict = {"source": indexed_source}
                excerpts = extract_reference_excerpts(valid_refs, indexed_dict)

                references_info = {
                    "enabled": True,
                    "count": len(valid_refs),
                    "valid": valid,
                    "errors": errors,
                    "refs": excerpts
                }
        else:
            log(f"✗ ({result.error})")
            # Fallback: combiner les extractions
            self.session.synthese_finale = "SYNTHÈSE (mode dégradé):\n" + "\n".join([
                r["content"][:200] for r in round1_results if r["success"]
            ])
        
        # Finalisation
        self.session.duration_total = time.time() - start_time
        self.session.timestamp_end = datetime.now().isoformat()
        
        # Sauvegarder les métriques de retry (PRD-004)
        if self.retry_metrics:
            self.session.retry_metrics = self.retry_metrics.to_dict()

            # Afficher résumé des retries si pertinent
            summary = self.retry_metrics.summary()
            if "retry" in summary.lower() and "aucun" not in summary.lower():
                log(f"\n{summary}")

        # PRD-029 Story 29.1: Sauvegarder les métriques de cache
        if self.cache and self.cache.enabled:
            self.session.cache_metrics = self.cache.stats()
            if self.session.cache_metrics.get("hits", 0) > 0:
                log(f"\n⚡ Cache: {self.session.cache_metrics['hits']} hits, "
                    f"{self.session.cache_metrics['latency_saved_s']}s économisés")

        log(f"\n═══ TOTAL: {self.session.duration_total:.1f}s ═══")

        # PRD-017: Ajouter les infos personas à la session
        if personas_info:
            self.session.personas_info = personas_info

        # PRD-019: Ajouter les infos références à la session
        if references_info:
            self.session.references_info = references_info

        # Afficher le résumé pédagogique (PRD-005)
        if self.pedagogy and self.pedagogy.enabled:
            self.pedagogy.explain(PedagogicStep.SUMMARY)
            stats = self.pedagogy.summary_stats(self.session)
            if stats:
                log(stats)

        return self.session
    
    # ══════════════════════════════════════════════════════════════════════════
    # MÉTHODE RUN SYNCHRONE (compatibilité avec les tests)
    # ══════════════════════════════════════════════════════════════════════════
    
    def run(
        self,
        source_text: str,
        cadrage: Dict[str, str],
        mode: Mode = Mode.STANDARD,
        personas: str = None
    ) -> SyntheseSession:
        """
        Exécute le processus de synthèse.

        Version wrapper qui appelle run_async via asyncio.run() si possible,
        sinon utilise une version synchrone pour compatibilité.
        """
        try:
            # Essayer d'obtenir un event loop existant
            loop = asyncio.get_running_loop()
            # Si on est déjà dans un event loop, créer une tâche
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(asyncio.run, self.run_async(source_text, cadrage, mode, personas=personas))
                return future.result()
        except RuntimeError:
            # Pas d'event loop en cours, utiliser asyncio.run directement
            return asyncio.run(self.run_async(source_text, cadrage, mode, personas=personas))

    # ══════════════════════════════════════════════════════════════════════════
    # MODE DIRECT (PRD-016)
    # ══════════════════════════════════════════════════════════════════════════

    def _build_direct_prompt(self, source_text: str, cadrage: Dict[str, str]) -> str:
        """Construit le prompt pour le mode direct (PRD-016)."""
        return f"""Tu es un expert en synthèse de documents.

CADRAGE:
- Destinataire: {cadrage.get('destinataire', 'non spécifié')}
- Finalité: {cadrage.get('finalite', 'non spécifiée')}
- Longueur: {cadrage.get('longueur', 'non spécifiée')}
- Ton: {cadrage.get('ton', 'non spécifié')}
- Niveau: {cadrage.get('niveau', 'non spécifié')}

TEXTE À SYNTHÉTISER:
---
{source_text}
---

CONSIGNES:
Produis une synthèse fidèle au texte source, adaptée au cadrage ci-dessus.
Évite les glissements de sens et les ajouts non présents dans le texte original.
Structure ta réponse de manière claire et exploitable par le destinataire.
"""

    async def run_direct_async(
        self,
        source_text: str,
        model: str,
        cadrage: Dict[str, str]
    ) -> SyntheseSession:
        """
        Exécute une synthèse directe avec un seul modèle (PRD-016).

        Args:
            source_text: Texte à synthétiser
            model: Modèle à utiliser (claude, gemini, codex)
            cadrage: Paramètres de cadrage

        Returns:
            SyntheseSession avec la synthèse directe
        """
        start_time = time.time()
        timestamp_start = datetime.now().isoformat()

        # Vérifier disponibilité du modèle
        if model not in self.available_models:
            raise ValueError(
                f"Modèle '{model}' non disponible. "
                f"Disponibles: {', '.join(self.available_models)}"
            )

        # PRD-008 Story 8.4: Sanitiser les entrées
        if SANITIZE_MODULE_AVAILABLE:
            clean_text, clean_cadrage, _, warnings = validate_and_sanitize(
                source_text, cadrage, strict=False
            )
            if warnings:
                for w in warnings:
                    log(f"⚠ {w}")
            source_text = clean_text
            cadrage = clean_cadrage

        # PRD-029: Stocker cadrage pour clé de cache
        self.cadrage = cadrage

        # Construire le prompt
        prompt = self._build_direct_prompt(source_text, cadrage)

        # Exécuter
        log(f"\n─── MODE DIRECT : {model} ───")
        result = await self._call_model_smart(
            model=model,
            prompt=prompt,
            timeout=self.config.get("timeout", 300),
            role="direct"
        )

        if not result.success:
            raise RuntimeError(f"Échec du modèle {model}: {result.error}")

        log(f"└─ {model}... ✓ ({result.duration:.1f}s)")

        duration = time.time() - start_time
        log(f"\n═══ TOTAL: {duration:.1f}s ═══\n")

        return SyntheseSession(
            session_id=self._generate_session_id(),
            query="synthèse directe",
            source_text=source_text[:500] + "..." if len(source_text) > 500 else source_text,
            mode="direct",
            cadrage=cadrage,
            rounds=[],  # Pas de rounds en mode direct
            synthese_finale=result.content,
            convergence=1.0,  # Pas de convergence en mono-modèle
            duration_total=duration,
            timestamp_start=timestamp_start,
            timestamp_end=datetime.now().isoformat(),
            models_used=[model]
        )

    def run_direct(
        self,
        source_text: str,
        model: str,
        cadrage: Dict[str, str]
    ) -> SyntheseSession:
        """
        Version synchrone de run_direct_async (PRD-016).
        """
        try:
            loop = asyncio.get_running_loop()
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(
                    asyncio.run,
                    self.run_direct_async(source_text, model, cadrage)
                )
                return future.result()
        except RuntimeError:
            return asyncio.run(self.run_direct_async(source_text, model, cadrage))

    def save_trail(self, session: SyntheseSession) -> Optional[str]:
        """Sauvegarde le trail de la session"""
        if not self.config["enable_trail"]:
            return None
        
        trail_dir = Path(self.config["trail_dir"])
        
        try:
            trail_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            print(f"Erreur création répertoire trail: {e}")
            return None
        
        trail_file = trail_dir / f"{session.session_id}.json"
        
        try:
            with open(trail_file, "w", encoding="utf-8") as f:
                json.dump(asdict(session), f, ensure_ascii=False, indent=2)
            return str(trail_file)
        except Exception as e:
            log(f"Erreur sauvegarde trail: {e}")
            return None

    def save_markdown(self, session: SyntheseSession, with_frontmatter: bool = False) -> Optional[str]:
        """
        Sauvegarde automatique de la synthèse en Markdown (PRD-014).

        PRD-022: Support optionnel du frontmatter YAML.

        Args:
            session: Session de synthese
            with_frontmatter: Ajouter le frontmatter YAML

        Returns:
            Chemin du fichier créé ou None si erreur
        """
        output_dir = Path("output")

        try:
            output_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            log(f"Erreur création répertoire output: {e}")
            return None

        md_file = output_dir / f"{session.session_id}.md"

        try:
            content = ""
            # PRD-022: Frontmatter optionnel
            if with_frontmatter:
                content += generate_frontmatter(session)
            content += format_output_markdown(session)
            with open(md_file, "w", encoding="utf-8") as f:
                f.write(content)
            return str(md_file)
        except Exception as e:
            log(f"Erreur sauvegarde markdown: {e}")
            return None

    def load_session(self, trail_path: Path) -> SyntheseSession:
        """
        Charge une session depuis un trail JSON.

        PRD-025: Reprise de session interrompue.

        Args:
            trail_path: Chemin vers le fichier trail

        Returns:
            Session reconstruite
        """
        with open(trail_path, encoding="utf-8") as f:
            data = json.load(f)

        return SyntheseSession(
            session_id=data["session_id"],
            query=data.get("query", "resume"),
            source_text=data["source_text"],
            mode=data["mode"],
            cadrage=data["cadrage"],
            rounds=data.get("rounds", []),
            synthese_finale=data.get("synthese_finale", ""),
            convergence=data.get("convergence", 0),
            duration_total=data.get("duration_total", 0),
            timestamp_start=data.get("timestamp_start", ""),
            timestamp_end=data.get("timestamp_end", ""),
            models_used=data.get("models_used", []),
            retry_metrics=data.get("retry_metrics", {}),
            personas_info=data.get("personas_info", {}),
            sources_info=data.get("sources_info", {}),
            references_info=data.get("references_info", {})
        )

    def get_next_stage(self, session: SyntheseSession) -> Optional[str]:
        """
        Détermine le prochain stage à exécuter.

        PRD-025: Identification du point de reprise.

        Args:
            session: Session chargée

        Returns:
            'extraction', 'critique', 'synthese', ou None si complet
        """
        completed_stages = {r["stage"] for r in session.rounds}

        if "extraction" not in completed_stages:
            return "extraction"

        mode = Mode[session.mode.upper()] if isinstance(session.mode, str) else session.mode

        if mode in [Mode.STANDARD, Mode.CRITIQUE]:
            if "critique" not in completed_stages:
                return "critique"

        if not session.synthese_finale:
            return "synthese"

        return None


# ══════════════════════════════════════════════════════════════════════════════
# INTERFACE UTILISATEUR
# ══════════════════════════════════════════════════════════════════════════════

def prompt_cadrage() -> Dict[str, str]:
    """Dialogue de cadrage interactif"""
    print("\n╔══════════════════════════════════════════════════════════════╗")
    print("║              CADRAGE DE LA SYNTHÈSE                          ║")
    print("╚══════════════════════════════════════════════════════════════╝\n")
    
    cadrage = {}
    
    cadrage["destinataire"] = input("À qui s'adresse cette synthèse ? [général] ").strip() or "public général"
    cadrage["finalite"] = input("Pour quoi faire ? (décision/information/argumentation) [information] ").strip() or "information"
    cadrage["longueur"] = input("Quelle longueur ? (ex: 10 lignes, 200 mots) [10-15 lignes] ").strip() or "10-15 lignes"
    cadrage["ton"] = input("Quel ton ? (formel/accessible/technique) [accessible] ").strip() or "accessible"
    cadrage["niveau"] = input("Quel niveau d'expertise ? (débutant/intermédiaire/expert) [intermédiaire] ").strip() or "intermédiaire"
    
    return cadrage


def format_output(session: SyntheseSession) -> str:
    """Formate la sortie pour l'utilisateur"""
    output = []
    output.append("\n" + "═" * 64)
    output.append("                    SYNTHÈSE CO-FABRIQUÉE")
    output.append("═" * 64)
    output.append(f"\nSession: {session.session_id}")
    output.append(f"Modèles: {', '.join(session.models_used)}")

    # PRD-018: Afficher les sources si multiples
    if session.sources_info and session.sources_info.get("count", 0) > 1:
        count = session.sources_info["count"]
        mode = session.sources_info.get("mode", "merge")
        output.append(f"Sources: {count} (mode {mode})")

    # PRD-017: Afficher les personas si disponibles
    if session.personas_info and session.personas_info.get("personas"):
        personas = session.personas_info["personas"]
        domain = session.personas_info.get("domain")
        if domain:
            output.append(f"Domaine: {domain}")
        output.append("Personas: " + ", ".join([p["name"] for p in personas]))

    output.append(f"Convergence: {session.convergence:.0%}")
    output.append(f"Durée: {session.duration_total:.1f}s")
    output.append("\n" + "─" * 64)
    output.append("\n" + session.synthese_finale)
    output.append("\n" + "═" * 64)

    return "\n".join(output)


def format_output_markdown(session: SyntheseSession) -> str:
    """Formate la sortie en Markdown (PRD-014)."""

    output = []

    # Titre
    output.append("# Synthèse Co-fabriquée\n")

    # Métadonnées en tableau
    output.append("## Métadonnées\n")
    output.append("| Champ | Valeur |")
    output.append("|-------|--------|")
    output.append(f"| Session | `{session.session_id}` |")
    output.append(f"| Mode | {session.mode} |")
    output.append(f"| Modèles | {', '.join(session.models_used)} |")
    output.append(f"| Convergence | {session.convergence}% |")
    output.append(f"| Durée | {session.duration_total:.1f}s |")
    output.append("")

    # PRD-018: Sources si multiples
    if session.sources_info and session.sources_info.get("count", 0) > 1:
        output.append("## Sources\n")
        mode = session.sources_info.get("mode", "merge")
        output.append(f"**Mode** : {mode}\n")
        for src in session.sources_info.get("sources", []):
            output.append(f"- **{src['name']}** ({src['type']})")
        output.append("")

    # PRD-017: Personas si disponibles
    if session.personas_info and session.personas_info.get("personas"):
        output.append("## Personas\n")
        domain = session.personas_info.get("domain")
        if domain:
            output.append(f"**Domaine détecté** : {domain}\n")
        for p in session.personas_info["personas"]:
            output.append(f"- **{p['name']}** : {p['focus']}")
        output.append("")

    # Cadrage
    output.append("## Cadrage\n")
    for key, value in session.cadrage.items():
        output.append(f"- **{key.capitalize()}** : {value}")
    output.append("")

    # Synthèse
    output.append("## Synthèse\n")
    output.append(session.synthese_finale)
    output.append("")

    # Footer
    output.append("---")
    output.append(f"*Trail : `synthese_trails/{session.session_id}.json`*")

    return "\n".join(output)


# ══════════════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE ASYNC (Story 2.4)
# ══════════════════════════════════════════════════════════════════════════════

async def async_main(args: argparse.Namespace) -> int:
    """
    Point d'entrée asynchrone principal.

    Story 2.4 du PRD-002: Permet l'exécution async depuis le CLI.
    PRD-003: Intégration de la configuration dynamique.
    """

    # PRD-025: Reprise de session
    if args.resume:
        council = SyntheseCouncil()
        trail_path = Path("synthese_trails") / f"{args.resume}.json"

        if not trail_path.exists():
            # Essayer avec ID partiel
            matches = list(Path("synthese_trails").glob(f"*{args.resume}*.json"))
            if len(matches) == 1:
                trail_path = matches[0]
            elif len(matches) > 1:
                print(f"❌ Plusieurs sessions correspondent à '{args.resume}':")
                for m in matches[:5]:
                    print(f"   {m.stem}")
                return 1
            else:
                print(f"❌ Session introuvable: {args.resume}")
                return 1

        log(f"🔄 Reprise de session: {trail_path.stem}")
        session = council.load_session(trail_path)

        next_stage = council.get_next_stage(session)
        if next_stage is None:
            log("✓ Session déjà complète")
        else:
            log(f"   Prochain stage: {next_stage}")
            # TODO: Implémenter resume_session() pour continuer l'exécution

        # Afficher le résultat existant
        if args.json:
            print(json.dumps(asdict(session), ensure_ascii=False, indent=2))
        elif args.markdown:
            print(format_output_markdown(session))
        else:
            print(format_output(session))

        return 0

    # PRD-018: Récupération des sources (fichiers, URLs, texte)
    sources = []
    source_mode = "single"  # Par défaut
    urls_arg = getattr(args, 'url', None)

    if args.file or urls_arg or args.text:
        # Charger toutes les sources
        sources = load_sources(
            files=args.file,
            urls=urls_arg,
            text=args.text
        )

        if not sources:
            print("Erreur: aucune source n'a pu être chargée")
            return 1

        # Déterminer le mode multi-sources (compare par défaut si plusieurs sources)
        compare_mode = getattr(args, 'compare', False)
        merge_mode = getattr(args, 'merge', False)

        if compare_mode:
            source_mode = "compare"
        elif merge_mode:
            source_mode = "merge"
        else:
            # Par défaut: compare si plusieurs sources, sinon merge
            source_mode = "compare" if len(sources) > 1 else "merge"

        # Formater les sources pour le prompt
        source_text = format_sources_for_prompt(sources, mode=source_mode)

        # Log multi-sources
        if len(sources) > 1:
            log(f"📚 {len(sources)} sources chargées (mode {source_mode})")
            for src in sources:
                log(f"   • {src.name} ({src.source_type})")
    else:
        # Mode interactif
        print("Entrez le texte à synthétiser (terminez par une ligne vide):")
        lines = []
        while True:
            line = input()
            if line == "":
                break
            lines.append(line)
        source_text = "\n".join(lines)
        sources = [Source(name="Texte interactif", content=source_text, source_type="text")]

    if not source_text or not source_text.strip():
        print("Erreur: aucun texte fourni")
        return 1
    
    # Construction des overrides CLI (PRD-003)
    cli_overrides = {}
    if hasattr(args, 'timeout') and args.timeout is not None:
        cli_overrides["timeout"] = args.timeout
    if hasattr(args, 'max_rounds') and args.max_rounds is not None:
        cli_overrides["max_rounds"] = args.max_rounds
    if args.no_trail:
        cli_overrides["enable_trail"] = False
    
    # Chemin de config personnalisé
    config_path = getattr(args, 'config', None)
    
    # Créer le council avec configuration (PRD-003)
    council = SyntheseCouncil(config=cli_overrides, config_path=config_path)
    
    # PRD-020: Cadrage via template ou CLI
    template_name = getattr(args, 'template', None)
    has_cadrage_args = any([args.destinataire, args.finalite, args.longueur, args.ton, args.niveau])

    if template_name or has_cadrage_args:
        # Utiliser resolve_cadrage (template + overrides CLI)
        if CONFIG_MODULE_AVAILABLE:
            default_cadrage = get_default_cadrage(council.config)
        else:
            default_cadrage = {
                "destinataire": "public général",
                "finalite": "information",
                "longueur": "10-15 lignes",
                "ton": "accessible",
                "niveau": "intermédiaire"
            }

        cadrage = resolve_cadrage(args, default_cadrage)

        if template_name:
            log(f"📋 Template: {template_name}")
    else:
        cadrage = prompt_cadrage()
    
    # Setup et exécution
    if not council.setup():
        return 1
    
    mode = Mode(args.mode)

    # Option pédagogique interactive (PRD-005)
    interactive_pedagogy = getattr(args, 'interactive_pedagogy', False)

    # PRD-017: Récupérer l'argument personas
    personas = getattr(args, 'personas', None)

    # PRD-019: Récupérer l'argument refs
    with_refs = getattr(args, 'refs', False)

    # Appel asynchrone (PRD-016: mode direct ou délibération)
    if mode == Mode.DIRECT:
        # Mode direct : un seul modèle, sans délibération
        session = await council.run_direct_async(source_text, args.model, cadrage)
    else:
        # Modes délibération : standard, rapide, critique, pédagogique
        session = await council.run_async(
            source_text, cadrage, mode, interactive_pedagogy,
            personas=personas, with_refs=with_refs
        )

    # PRD-018: Ajouter les informations de sources
    if sources and len(sources) > 0:
        session.sources_info = {
            "count": len(sources),
            "mode": source_mode,
            "sources": [
                {
                    "name": src.name,
                    "type": src.source_type,
                    "metadata": src.metadata
                }
                for src in sources
            ]
        }

    # PRD-021: Mode confirmation (boucle de validation)
    confirm_mode = getattr(args, 'confirm', False)

    # Ignorer --confirm si non-interactif (pipe, redirection)
    if confirm_mode and not sys.stdin.isatty():
        log("⚠️  Mode --confirm ignoré (entrée non-interactive)")
        confirm_mode = False

    if confirm_mode:
        regeneration_count = 0
        while True:
            # Afficher le brouillon
            display_draft(session)

            # Demander confirmation
            choice = prompt_confirmation()

            if choice == 'export':
                # Continuer vers l'export normal
                break

            elif choice == 'regenerate':
                regeneration_count += 1
                log(f"\n🔄 Régénération #{regeneration_count}...")
                if mode == Mode.DIRECT:
                    session = await council.run_direct_async(source_text, args.model, cadrage)
                else:
                    session = await council.run_async(
                        source_text, cadrage, mode, interactive_pedagogy,
                        personas=personas, with_refs=with_refs
                    )

            elif choice == 'modify':
                cadrage = prompt_modify_cadrage(cadrage)
                regeneration_count += 1
                log(f"\n🔄 Régénération #{regeneration_count} avec nouveau cadrage...")
                if mode == Mode.DIRECT:
                    session = await council.run_direct_async(source_text, args.model, cadrage)
                else:
                    session = await council.run_async(
                        source_text, cadrage, mode, interactive_pedagogy,
                        personas=personas, with_refs=with_refs
                    )

            elif choice == 'cancel':
                print("\n❌ Annulé. Aucun fichier sauvegardé.")
                return 0

    # Sauvegarde du trail JSON
    if council.config.get("enable_trail", True):
        trail_path = council.save_trail(session)
        if trail_path:
            log(f"\nTrail sauvegardé: {trail_path}")

    # Sauvegarde automatique du Markdown (PRD-014)
    # PRD-022: Support frontmatter
    with_frontmatter = getattr(args, 'frontmatter', False)
    md_path = council.save_markdown(session, with_frontmatter=with_frontmatter)
    if md_path:
        log(f"Synthèse Markdown: {md_path}")

    # Sortie
    if args.json:
        print(json.dumps(asdict(session), ensure_ascii=False, indent=2))
    elif args.markdown:
        # PRD-022: Support frontmatter sur stdout
        if with_frontmatter:
            print(generate_frontmatter(session) + format_output_markdown(session))
        else:
            print(format_output_markdown(session))
    else:
        print(format_output(session))

    # PRD-026: Affichage des métriques détaillées
    if getattr(args, 'metrics', False):
        print("\n═══ MÉTRIQUES ═══")
        print(f"Durée totale: {session.duration_total:.1f}s")
        print(f"Modèles: {', '.join(session.models_used)}")
        print(f"Rounds: {len(session.rounds)}")
        if session.retry_metrics:
            total_retries = sum(
                m.get('total_attempts', 1) - 1
                for m in session.retry_metrics.values()
            )
            if total_retries > 0:
                print(f"Retries: {total_retries}")
        print(f"Convergence: {session.convergence:.0%}")

    return 0


def main():
    """Point d'entrée synchrone (wrapper pour async_main)"""
    
    parser = argparse.ArgumentParser(
        description="Synthèse co-fabriquée par conseil de LLMs (v2.2 - mode pédagogique)"
    )
    
    # ═══════════════════════════════════════════════════════════════════
    # COMMANDES DE BASE
    # ═══════════════════════════════════════════════════════════════════
    
    parser.add_argument(
        "--check",
        action="store_true",
        help="Vérifier les CLI disponibles"
    )
    parser.add_argument(
        "--file", "-f",
        type=str,
        nargs="+",
        help="Fichier(s) source à synthétiser (txt, pdf, docx)"
    )
    parser.add_argument(
        "--text", "-t",
        type=str,
        help="Texte source direct"
    )
    parser.add_argument(
        "--url",
        type=str,
        nargs="+",
        help="URL(s) de pages web à synthétiser"
    )
    parser.add_argument(
        "--compare",
        action="store_true",
        help="Mode comparaison : analyse les différences entre sources"
    )
    parser.add_argument(
        "--merge",
        action="store_true",
        help="Mode fusion : combine les sources en une synthèse unifiée (défaut)"
    )
    parser.add_argument(
        "--mode", "-m",
        choices=["standard", "rapide", "critique", "pedagogique", "direct"],
        default="standard",
        help="Mode de synthèse"
    )
    parser.add_argument(
        "--direct",
        action="store_true",
        help="Mode direct : un seul modèle, sans délibération (alias pour --mode direct)"
    )
    parser.add_argument(
        "--model",
        type=str,
        choices=["claude", "gemini", "codex"],
        help="Modèle à utiliser (requis avec --direct ou --mode direct)"
    )
    parser.add_argument(
        "--personas",
        type=str,
        help="Personas experts: 'auto', preset (tech/science/legal/business/medical), ou liste personnalisée"
    )
    parser.add_argument(
        "--refs", "--references",
        action="store_true",
        help="Inclure des références [§N] aux passages sources (PRD-019)"
    )
    parser.add_argument(
        "--template", "-T",
        type=str,
        choices=list(CADRAGE_TEMPLATES.keys()),
        help="Template de cadrage prédéfini (PRD-020)"
    )
    parser.add_argument(
        "--list-templates",
        action="store_true",
        help="Afficher les templates de cadrage disponibles"
    )
    parser.add_argument(
        "--confirm", "-C",
        action="store_true",
        help="Demander confirmation avant l'export final (PRD-021)"
    )
    parser.add_argument(
        "--frontmatter", "-F",
        action="store_true",
        help="Ajouter un frontmatter YAML aux fichiers Markdown (PRD-022)"
    )

    # ═══════════════════════════════════════════════════════════════════
    # CACHE (PRD-024)
    # ═══════════════════════════════════════════════════════════════════

    parser.add_argument(
        "--no-cache",
        action="store_true",
        help="Désactiver le cache des réponses (PRD-024)"
    )
    parser.add_argument(
        "--clear-cache",
        action="store_true",
        help="Vider le cache des réponses"
    )
    parser.add_argument(
        "--cache-stats",
        action="store_true",
        help="Afficher les statistiques du cache"
    )

    # ═══════════════════════════════════════════════════════════════════
    # REPRISE DE SESSION (PRD-025)
    # ═══════════════════════════════════════════════════════════════════

    parser.add_argument(
        "--resume", "-R",
        type=str,
        metavar="SESSION_ID",
        help="Reprendre une session interrompue (PRD-025)"
    )
    parser.add_argument(
        "--list-sessions",
        action="store_true",
        help="Lister les sessions disponibles pour reprise"
    )

    # ═══════════════════════════════════════════════════════════════════
    # MÉTRIQUES (PRD-026)
    # ═══════════════════════════════════════════════════════════════════

    parser.add_argument(
        "--metrics",
        action="store_true",
        help="Afficher les métriques détaillées (PRD-026)"
    )

    # ═══════════════════════════════════════════════════════════════════
    # ARGUMENTS DE CADRAGE
    # ═══════════════════════════════════════════════════════════════════
    
    parser.add_argument(
        "--destinataire",
        type=str,
        help="Destinataire de la synthèse"
    )
    parser.add_argument(
        "--finalite",
        type=str,
        help="Finalité de la synthèse"
    )
    parser.add_argument(
        "--longueur",
        type=str,
        help="Longueur souhaitée"
    )
    parser.add_argument(
        "--ton",
        type=str,
        help="Ton de la synthèse"
    )
    parser.add_argument(
        "--niveau",
        type=str,
        help="Niveau d'expertise"
    )
    
    # ═══════════════════════════════════════════════════════════════════
    # ARGUMENTS DE CONFIGURATION (PRD-003)
    # ═══════════════════════════════════════════════════════════════════
    
    parser.add_argument(
        "--config", "-c",
        type=str,
        help="Chemin vers un fichier de configuration YAML"
    )
    parser.add_argument(
        "--validate-config",
        action="store_true",
        help="Valider le fichier de configuration"
    )
    parser.add_argument(
        "--show-config",
        action="store_true",
        help="Afficher la configuration effective"
    )
    parser.add_argument(
        "--init-config",
        action="store_true",
        help="Créer un fichier de configuration template"
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Forcer l'écrasement (avec --init-config)"
    )
    
    # ═══════════════════════════════════════════════════════════════════
    # AUTRES ARGUMENTS
    # ═══════════════════════════════════════════════════════════════════
    
    parser.add_argument(
        "--no-trail",
        action="store_true",
        help="Désactiver la sauvegarde du trail"
    )
    parser.add_argument(
        "--json",
        action="store_true",
        help="Sortie JSON"
    )
    parser.add_argument(
        "--markdown", "--md",
        action="store_true",
        help="Sortie Markdown formatée"
    )
    parser.add_argument(
        "--timeout",
        type=int,
        help="Timeout par modèle en secondes (override)"
    )
    parser.add_argument(
        "--max-rounds",
        type=int,
        help="Nombre de rounds (override)"
    )
    parser.add_argument(
        "--interactive-pedagogy",
        action="store_true",
        help="Mode pédagogique interactif (pause après chaque explication)"
    )
    
    args = parser.parse_args()

    # ═══════════════════════════════════════════════════════════════════
    # VALIDATION MODE DIRECT (PRD-016)
    # ═══════════════════════════════════════════════════════════════════

    # --direct est un alias pour --mode direct
    if args.direct:
        args.mode = "direct"

    # Vérifier que --model est fourni pour le mode direct
    if args.mode == "direct" and not args.model:
        parser.error("--direct (ou --mode direct) nécessite --model (claude, gemini, codex)")

    # ═══════════════════════════════════════════════════════════════════
    # COMMANDES DE CONFIGURATION (PRD-003 Story 3.5)
    # ═══════════════════════════════════════════════════════════════════

    if args.init_config:
        # Créer un fichier de configuration template
        if not CONFIG_MODULE_AVAILABLE:
            print("Erreur: Module de configuration non disponible.")
            print("Installez PyYAML: pip install pyyaml")
            sys.exit(1)
        
        try:
            path = ConfigLoader.init_config(force=args.force)
            print(f"✓ Fichier de configuration créé: {path}")
            sys.exit(0)
        except ConfigError as e:
            print(f"✗ {e}")
            sys.exit(1)
    
    if args.validate_config:
        # Valider le fichier de configuration
        if not CONFIG_MODULE_AVAILABLE:
            print("Erreur: Module de configuration non disponible.")
            sys.exit(1)
        
        try:
            config = ConfigLoader.load(path=args.config)
            print("╔══════════════════════════════════════════════════════════════╗")
            print("║           VALIDATION DE LA CONFIGURATION                     ║")
            print("╠══════════════════════════════════════════════════════════════╣")
            
            # Afficher les sources
            sources = config.get("_sources", ["default"])
            for source in sources:
                print(f"║  ✓ Source chargée: {source:<46} ║")
            
            print("╠══════════════════════════════════════════════════════════════╣")
            print("║  Configuration valide ✓                                      ║")
            print("╚══════════════════════════════════════════════════════════════╝")
            sys.exit(0)
        except ConfigError as e:
            print("╔══════════════════════════════════════════════════════════════╗")
            print("║           VALIDATION DE LA CONFIGURATION                     ║")
            print("╠══════════════════════════════════════════════════════════════╣")
            print("║  Configuration INVALIDE ✗                                    ║")
            print("╠══════════════════════════════════════════════════════════════╣")
            for err in e.errors:
                print(f"║  • {err:<58} ║")
            print("╚══════════════════════════════════════════════════════════════╝")
            sys.exit(1)
    
    if args.show_config:
        # Afficher la configuration effective
        if not CONFIG_MODULE_AVAILABLE:
            print("Erreur: Module de configuration non disponible.")
            sys.exit(1)
        
        try:
            # Construire les overrides CLI
            cli_overrides = {}
            if args.timeout:
                cli_overrides["timeout"] = args.timeout
            if args.max_rounds:
                cli_overrides["max_rounds"] = args.max_rounds
            if args.no_trail:
                cli_overrides["enable_trail"] = False
            
            config = ConfigLoader.load(path=args.config, cli_overrides=cli_overrides)
            
            print("╔══════════════════════════════════════════════════════════════╗")
            print("║           CONFIGURATION EFFECTIVE                            ║")
            print("╠══════════════════════════════════════════════════════════════╣")
            
            # Afficher les sources
            sources = config.get("_sources", ["default"])
            print(f"║  Sources: {' → '.join(sources):<51} ║")
            print("╠══════════════════════════════════════════════════════════════╣")
            
            # Afficher les chemins
            paths = ConfigLoader.get_config_paths()
            for name, (path, exists) in paths.items():
                status = "✓" if exists else "✗"
                print(f"║  {status} {name}: {str(path)[:50]:<50} ║")
            
            print("╚══════════════════════════════════════════════════════════════╝")
            print("\n" + ConfigLoader.show_config(config))
            sys.exit(0)
        except ConfigError as e:
            print(f"Erreur: {e}")
            sys.exit(1)
    
    # ═══════════════════════════════════════════════════════════════════
    # COMMANDE DE VÉRIFICATION CLI
    # ═══════════════════════════════════════════════════════════════════
    
    if args.check:
        results = check_all_clis()
        print("\n╔══════════════════════════════════════════════════════════════╗")
        print("║           VÉRIFICATION DES CLI                               ║")
        print("╠══════════════════════════════════════════════════════════════╣")
        for model, (available, info) in results.items():
            status = "✓" if available else "✗"
            print(f"║  {status} {model:10} │ {info[:45]:<45} ║")
        print("╚══════════════════════════════════════════════════════════════╝")
        
        available_count = sum(1 for _, (a, _) in results.items() if a)
        sys.exit(0 if available_count > 0 else 1)

    # PRD-020: Afficher les templates disponibles
    if args.list_templates:
        list_templates()
        sys.exit(0)

    # PRD-024: Gestion du cache
    if args.cache_stats:
        cache = ResponseCache()
        stats = cache.stats()
        print(f"📦 Cache: {stats['entries']} entrées ({stats['size_mb']} MB)")
        if stats.get('cache_dir'):
            print(f"   Répertoire: {stats['cache_dir']}")
        sys.exit(0)

    if args.clear_cache:
        cache = ResponseCache()
        count = cache.clear()
        print(f"🗑️  Cache vidé: {count} entrées supprimées")
        sys.exit(0)

    # PRD-025: Lister les sessions disponibles
    if args.list_sessions:
        trail_dir = Path("synthese_trails")
        if not trail_dir.exists():
            print("Aucune session disponible")
            sys.exit(0)
        print("📋 Sessions disponibles:\n")
        for trail in sorted(trail_dir.glob("*.json"), reverse=True)[:10]:
            try:
                with open(trail, encoding="utf-8") as f:
                    data = json.load(f)
                completed = "✓" if data.get("synthese_finale") else "⏸️"
                rounds = len(data.get("rounds", []))
                print(f"  {completed} {data['session_id']}")
                print(f"     Mode: {data.get('mode', '?')} | Rounds: {rounds}")
            except Exception:
                continue
        sys.exit(0)

    # ═══════════════════════════════════════════════════════════════════
    # EXÉCUTION PRINCIPALE
    # ═══════════════════════════════════════════════════════════════════

    # PRD-014: Rediriger logs vers stderr si sortie structurée (json/markdown)
    global _log_file
    if args.json or args.markdown:
        _log_file = sys.stderr
        _printer.set_file(sys.stderr)

    # Compatible Windows/Linux/macOS
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

    exit_code = asyncio.run(async_main(args))
    sys.exit(exit_code)


if __name__ == "__main__":
    main()
