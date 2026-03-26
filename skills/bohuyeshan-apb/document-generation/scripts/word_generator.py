"""
Word Generator Tool
Converts HTML content into a Microsoft Word (.docx) document.
Supports headings, paragraphs, and lists.
"""

import datetime
import os
from pathlib import Path
from typing import Dict
import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordGenerator:
    def __init__(self, output_dir: str = "generated_docs/word"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _parse_html(self, html: str) -> BeautifulSoup:
        return BeautifulSoup(html, "html.parser")

    def generate_docx(self, html_content: str, filename: str = None) -> str:
        """
        Generates a DOCX file from the provided HTML content.
        
        Args:
            html_content: HTML string.
            filename: Output filename.
            
        Returns:
            Absolute path to the generated DOCX file.
        """
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
            filename = f"document_{timestamp}.docx"
            
        if not filename.endswith(".docx"):
            filename += ".docx"

        output_path = self.output_dir / filename
        
        document = Document()
        soup = self._parse_html(html_content)

        def parse_css(style_text: str) -> Dict[str, str]:
            res: Dict[str, str] = {}
            if not style_text:
                return res
            for part in style_text.split(";"):
                if ":" not in part:
                    continue
                k, v = part.split(":", 1)
                k = k.strip().lower()
                v = v.strip()
                if not k or not v:
                    continue
                res[k] = v
            return res

        def parse_color(raw: str):
            t = (raw or "").strip()
            if not t:
                return None
            if t.startswith("#") and len(t) in (4, 7):
                if len(t) == 4:
                    r = int(t[1] * 2, 16)
                    g = int(t[2] * 2, 16)
                    b = int(t[3] * 2, 16)
                else:
                    r = int(t[1:3], 16)
                    g = int(t[3:5], 16)
                    b = int(t[5:7], 16)
                return RGBColor(r, g, b)
            m = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", t, flags=re.IGNORECASE)
            if m:
                r = max(0, min(255, int(m.group(1))))
                g = max(0, min(255, int(m.group(2))))
                b = max(0, min(255, int(m.group(3))))
                return RGBColor(r, g, b)
            return None

        def parse_font_size(raw: str):
            t = (raw or "").strip().lower()
            if not t:
                return None
            m = re.match(r"(\d+(\.\d+)?)\s*(px|pt)?", t)
            if not m:
                return None
            val = float(m.group(1))
            unit = m.group(3) or "px"
            if unit == "px":
                return Pt(val * 0.75)
            return Pt(val)

        def merge_style(base: Dict[str, object], extra: Dict[str, object]) -> Dict[str, object]:
            merged = dict(base)
            for k, v in extra.items():
                if v is None:
                    continue
                merged[k] = v
            return merged

        def style_from_tag(tag: Tag) -> Dict[str, object]:
            s: Dict[str, object] = {}
            name = tag.name.lower()
            if name in ("b", "strong"):
                s["bold"] = True
            if name in ("i", "em"):
                s["italic"] = True
            if name == "u":
                s["underline"] = True

            css = parse_css(tag.get("style", ""))
            if "font-weight" in css:
                fw = css["font-weight"].strip().lower()
                if fw in ("bold", "bolder") or (fw.isdigit() and int(fw) >= 600):
                    s["bold"] = True
                elif fw.isdigit() and int(fw) <= 400:
                    s["bold"] = False
            if "font-style" in css:
                fs = css["font-style"].strip().lower()
                if fs == "italic":
                    s["italic"] = True
                elif fs == "normal":
                    s["italic"] = False
            if "text-decoration" in css:
                td = css["text-decoration"].strip().lower()
                if "underline" in td:
                    s["underline"] = True
            if "font-size" in css:
                s["font_size"] = parse_font_size(css.get("font-size"))
            if "font-family" in css:
                ff = css["font-family"].split(",")[0].strip().strip("\"' ")
                if ff:
                    s["font_name"] = ff
            if "color" in css:
                s["color"] = parse_color(css.get("color"))
            return s

        def apply_paragraph_style(paragraph, tag: Tag):
            css = parse_css(tag.get("style", ""))
            ta = (css.get("text-align") or "").strip().lower()
            if ta == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif ta == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif ta == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif ta == "left":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def add_runs(paragraph, node, cur_style: Dict[str, object]):
            if isinstance(node, NavigableString):
                t = str(node)
                if not t:
                    return
                run = paragraph.add_run(t)
                if "bold" in cur_style:
                    run.bold = bool(cur_style.get("bold"))
                if "italic" in cur_style:
                    run.italic = bool(cur_style.get("italic"))
                if "underline" in cur_style:
                    run.underline = bool(cur_style.get("underline"))
                if cur_style.get("font_name"):
                    run.font.name = str(cur_style["font_name"])
                if cur_style.get("font_size"):
                    run.font.size = cur_style["font_size"]
                if cur_style.get("color"):
                    run.font.color.rgb = cur_style["color"]
                return

            if not isinstance(node, Tag):
                return

            if node.name and node.name.lower() == "br":
                paragraph.add_run("\n")
                return

            next_style = merge_style(cur_style, style_from_tag(node))
            for child in node.children:
                add_runs(paragraph, child, next_style)

        blocks = soup.find_all(["h1", "h2", "h3", "h4", "p", "ul", "ol", "img", "table"])
        
        for element in blocks:
            # Skip elements that are nested inside other block-level elements we process
            # e.g. <p> inside <table>, or <li> inside <ul> (though ul logic handles li)
            # Actually find_all is flat list, but includes nested.
            # We specifically want to ignore p/img/ul/ol if they are inside a table, 
            # because we handle the table as a unit.
            if element.find_parent("table"):
                continue

            name = element.name.lower()

            if name == "table":
                rows = element.find_all("tr")
                if not rows: continue
                
                # Determine dimensions
                max_cols = 0
                for row in rows:
                    cols = row.find_all(["td", "th"])
                    max_cols = max(max_cols, len(cols))
                
                if max_cols == 0: continue

                # Add table
                try:
                    table = document.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Table Grid'
                except Exception:
                    # Fallback if style doesn't exist
                    table = document.add_table(rows=len(rows), cols=max_cols)

                for i, row in enumerate(rows):
                    cols = row.find_all(["td", "th"])
                    for j, col in enumerate(cols):
                        if j >= max_cols: break
                        cell = table.cell(i, j)
                        cell.text = col.get_text(strip=True)
                        # Bold header
                        if col.name == "th":
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                
                document.add_paragraph() # Add spacing after table
                continue

            if name == "img":
                src = element.get("src")
                if not src: continue
                
                try:
                    # Handle local paths
                    if os.path.exists(src):
                        document.add_picture(src, width=Pt(400)) # Default width
                    # Handle URLs (basic)
                    elif src.startswith("http"):
                        import requests
                        from io import BytesIO
                        resp = requests.get(src)
                        if resp.status_code == 200:
                            document.add_picture(BytesIO(resp.content), width=Pt(400))
                except Exception as e:
                    print(f"Error adding image {src}: {e}")
                    para = document.add_paragraph(f"[Image: {src}]")
                continue

            if name in ("ul", "ol"):
                for li in element.find_all("li", recursive=False):
                    style_name = "List Bullet" if name == "ul" else "List Number"
                    try:
                        para = document.add_paragraph(style=style_name)
                    except:
                        # Fallback if style missing
                        para = document.add_paragraph()
                        if name == "ul": para.style = "List Paragraph"
                    
                    apply_paragraph_style(para, li)
                    for child in li.children:
                        add_runs(para, child, {})
                continue

            if name.startswith("h"):
                text = element.get_text(" ", strip=True)
                if not text:
                    continue
                level = int(name[1])
                para = document.add_heading("", level=level)
                for child in element.children:
                    add_runs(para, child, {"bold": True})
                continue

            if name == "p":
                text = element.get_text(" ", strip=True)
                if not text:
                    continue
                para = document.add_paragraph()
                apply_paragraph_style(para, element)
                for child in element.children:
                    add_runs(para, child, {})
                continue
                
        document.save(str(output_path))
        return str(output_path.absolute())

if __name__ == "__main__":
    generator = WordGenerator()
    html = """
    <h1>Project Report</h1>
    <h2>1. Introduction</h2>
    <p>This is a generated report using N-T-AI Office Suite.</p>
    <h2>2. Features</h2>
    <ul>
        <li>Fast generation</li>
        <li>Clean formatting</li>
    </ul>
    """
    print(generator.generate_docx(html))
